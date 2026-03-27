#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Практическая работа №2: Анализ скорости загрузки сайта и рекомендации по ускорению.
Сохраняет отчёт в DOCX и JSON.
Использование: python speed_audit.py <URL> [--output OUTPUT_NAME]
"""

import sys
import requests
from urllib.parse import urljoin, urlparse
from bs4 import BeautifulSoup
import time
import json
import argparse
import re
from datetime import datetime

# Настройки
TIMEOUT = 10
USER_AGENT = "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36"
CHECK_RESOURCES_LIMIT = 50  # ограничение на количество проверяемых ресурсов

# Для корректного вывода в консоль (Windows)
if sys.platform == 'win32':
    sys.stdout.reconfigure(encoding='utf-8')

# ----------------------------------------------------------------------
# Вспомогательные функции
# ----------------------------------------------------------------------
def get_html(url):
    """Получить HTML страницы и метаинформацию о загрузке"""
    try:
        headers = {'User-Agent': USER_AGENT}
        # Измеряем TTFB (время до первого байта)
        start_ttfb = time.time()
        response = requests.get(url, headers=headers, timeout=TIMEOUT, allow_redirects=True, stream=True)
        ttfb = time.time() - start_ttfb
        # Читаем содержимое (полное время загрузки)
        start_full = time.time()
        content = response.content
        full_load_time = time.time() - start_full
        # Определяем кодировку
        if response.encoding is None or response.encoding.lower() == 'iso-8859-1':
            response.encoding = response.apparent_encoding or 'utf-8'
        html = response.text
        return html, response, ttfb, full_load_time
    except Exception as e:
        print(f"Ошибка загрузки {url}: {e}")
        return None, None, None, None

def extract_resources(html, base_url):
    """Извлекает из HTML ссылки на CSS, JS, изображения"""
    soup = BeautifulSoup(html, 'html.parser')
    resources = {'css': [], 'js': [], 'images': []}
    
    # CSS
    for link in soup.find_all('link', rel='stylesheet'):
        href = link.get('href')
        if href:
            resources['css'].append(urljoin(base_url, href))
    
    # JS
    for script in soup.find_all('script', src=True):
        src = script['src']
        if src:
            resources['js'].append(urljoin(base_url, src))
    
    # Изображения
    for img in soup.find_all('img', src=True):
        src = img['src']
        if src:
            resources['images'].append(urljoin(base_url, src))
    
    return resources

def measure_resources(resources, base_url):
    """Измеряет размеры и количество ресурсов (HEAD запросы)"""
    total_size = 0
    total_requests = 0
    details = {'css': {'count': 0, 'size': 0}, 'js': {'count': 0, 'size': 0}, 'images': {'count': 0, 'size': 0}}
    
    for res_type, urls in resources.items():
        count = 0
        size = 0
        for url in urls[:CHECK_RESOURCES_LIMIT]:
            try:
                resp = requests.head(url, timeout=5, allow_redirects=True)
                if resp.status_code == 200:
                    content_length = int(resp.headers.get('content-length', 0))
                    size += content_length
                    count += 1
            except Exception:
                pass
        details[res_type]['count'] = count
        details[res_type]['size'] = size
        total_requests += count
        total_size += size
    
    return total_requests, total_size, details

def check_gzip(headers):
    """Проверяет, включён ли GZIP"""
    content_encoding = headers.get('Content-Encoding', '')
    return 'gzip' in content_encoding.lower()

def check_caching(headers):
    """Проверяет наличие заголовков кэширования"""
    cache_control = headers.get('Cache-Control', '')
    expires = headers.get('Expires', '')
    return bool(cache_control or expires)

def check_minification(html, css_urls, js_urls):
    """Проверяет, используется ли минификация (по наличию .min в именах файлов)"""
    # HTML минификация оценивается по отсутствию лишних пробелов (приблизительно)
    html_size = len(html)
    html_minified = html_size < 50000  # грубо, если страница большая, возможно не минифицирована
    
    css_minified = any('.min.css' in url for url in css_urls)
    js_minified = any('.min.js' in url for url in js_urls)
    
    return {
        'html': html_minified,
        'css': css_minified,
        'js': js_minified
    }

def check_cdn(url):
    """Проверяет, используется ли CDN (по домену)"""
    domain = urlparse(url).netloc.lower()
    cdn_domains = ['cloudflare', 'cdn', 'akamai', 'fastly', 'maxcdn', 'stackpath']
    return any(cdn in domain for cdn in cdn_domains)

def check_https(url):
    """Проверяет, используется ли HTTPS"""
    return url.startswith('https://')

def count_redirects(url):
    """Подсчитывает количество редиректов (через requests с allow_redirects=True)"""
    try:
        response = requests.get(url, timeout=TIMEOUT, allow_redirects=False)
        redirect_count = 0
        while response.status_code in (301, 302, 303, 307, 308):
            redirect_count += 1
            if 'location' in response.headers:
                next_url = response.headers['location']
                response = requests.get(next_url, timeout=TIMEOUT, allow_redirects=False)
            else:
                break
        return redirect_count
    except Exception:
        return 0

def analyze_performance(url):
    """Основная функция анализа производительности"""
    print(f"Анализ сайта: {url}")
    report = {
        "url": url,
        "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "load_time": {},
        "resources": {},
        "headers": {},
        "optimization": {},
        "recommendations": []
    }
    
    # 1. Загружаем главную страницу
    html, response, ttfb, full_load_time = get_html(url)
    if not html:
        report["error"] = "Сайт недоступен"
        return report
    
    report["load_time"] = {
        "ttfb_sec": round(ttfb, 2),
        "full_load_sec": round(full_load_time, 2)
    }
    
    # 2. Извлекаем ресурсы
    resources = extract_resources(html, url)
    total_requests, total_size, details = measure_resources(resources, url)
    report["resources"] = {
        "total_requests": total_requests,
        "total_size_bytes": total_size,
        "total_size_mb": round(total_size / (1024*1024), 2),
        "by_type": details
    }
    
    # 3. Анализируем заголовки
    headers = response.headers
    report["headers"] = {
        "gzip_enabled": check_gzip(headers),
        "caching_enabled": check_caching(headers),
        "cdn_used": check_cdn(url),
        "https_used": check_https(url),
        "redirects_count": count_redirects(url)
    }
    
    # 4. Проверка минификации
    minification = check_minification(html, resources['css'], resources['js'])
    report["optimization"]["minification"] = minification
    
    # 5. Формируем рекомендации
    recs = []
    
    if report["load_time"]["full_load_sec"] > 3:
        recs.append(f"Страница загружается {report['load_time']['full_load_sec']} с (рекомендуется < 3 с).")
    if report["load_time"]["ttfb_sec"] > 0.5:
        recs.append(f"TTFB = {report['load_time']['ttfb_sec']} с (желательно < 0.5 с). Оптимизируйте работу сервера, включите кэширование.")
    
    if not report["headers"]["gzip_enabled"]:
        recs.append("Включите GZIP сжатие в .htaccess или настройках сервера.")
    if not report["headers"]["caching_enabled"]:
        recs.append("Настройте кэширование браузера (Expires, Cache-Control).")
    if not report["headers"]["cdn_used"]:
        recs.append("Подключите CDN (например, Cloudflare) для ускорения доставки контента.")
    if not report["headers"]["https_used"]:
        recs.append("Переведите сайт на HTTPS — это повысит безопасность и может ускорить загрузку за счёт HTTP/2.")
    if report["headers"]["redirects_count"] > 1:
        recs.append(f"Обнаружено {report['headers']['redirects_count']} редиректов. Уменьшите количество перенаправлений.")
    
    if report["resources"]["total_requests"] > 50:
        recs.append(f"Слишком много HTTP-запросов ({report['resources']['total_requests']}). Объедините CSS и JS файлы.")
    if report["resources"]["total_size_mb"] > 2:
        recs.append(f"Общий вес страницы {report['resources']['total_size_mb']} МБ. Оптимизируйте изображения и код.")
    
    if not minification['css'] or not minification['js']:
        recs.append("Минифицируйте CSS и JavaScript (удалите пробелы, комментарии). Используйте инструменты вроде CSSNano, UglifyJS.")
    if not minification['html']:
        recs.append("Минифицируйте HTML (можно плагинами CMS или через .htaccess).")
    
    # Дополнительные рекомендации по изображениям
    if details['images']['count'] > 10 and details['images']['size'] > 500000:
        recs.append(f"Изображения занимают {round(details['images']['size']/(1024*1024),2)} МБ. Сожмите их и используйте формат WebP.")
        recs.append("Внедрите ленивую загрузку (loading='lazy') для изображений и iframe.")
    
    if details['js']['count'] > 10:
        recs.append("Используйте атрибуты async/defer для JavaScript, чтобы не блокировать отрисовку.")
    
    # Рекомендация по хостингу (общая)
    recs.append("Оцените возможность перехода на более быстрый хостинг (VPS, NVMe SSD) или использования HTTP/2/3.")
    
    report["recommendations"] = recs
    
    return report

def save_json(report_dict, filename):
    with open(filename, 'w', encoding='utf-8') as f:
        json.dump(report_dict, f, ensure_ascii=False, indent=2)
    print(f"JSON отчёт сохранён в {filename}")

def save_docx(report_dict, filename):
    """Сохраняет отчёт в DOCX с помощью python-docx"""
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("Для сохранения DOCX установите python-docx: pip install python-docx")
        return
    
    doc = Document()
    heading = doc.add_heading("Отчёт по анализу скорости загрузки сайта", 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    # Информация о сайте
    doc.add_heading("Общая информация", level=1)
    doc.add_paragraph(f"URL: {report_dict['url']}")
    doc.add_paragraph(f"Дата анализа: {report_dict['timestamp']}")
    
    if "error" in report_dict:
        doc.add_paragraph(f"Ошибка: {report_dict['error']}")
        doc.save(filename)
        return
    
    # Время загрузки
    doc.add_heading("1. Время загрузки", level=1)
    doc.add_paragraph(f"TTFB (время до первого байта): {report_dict['load_time']['ttfb_sec']} с")
    doc.add_paragraph(f"Полное время загрузки страницы: {report_dict['load_time']['full_load_sec']} с")
    
    # Ресурсы
    doc.add_heading("2. Ресурсы страницы", level=1)
    doc.add_paragraph(f"Общее количество HTTP-запросов: {report_dict['resources']['total_requests']}")
    doc.add_paragraph(f"Общий вес страницы: {report_dict['resources']['total_size_mb']} МБ")
    doc.add_paragraph(f"Детализация по типам:")
    for typ, data in report_dict['resources']['by_type'].items():
        doc.add_paragraph(f"  {typ.upper()}: {data['count']} файлов, {round(data['size']/(1024*1024),2)} МБ")
    
    # Заголовки сервера
    doc.add_heading("3. Настройки сервера и оптимизация", level=1)
    h = report_dict['headers']
    doc.add_paragraph(f"GZIP сжатие: {'включено' if h['gzip_enabled'] else 'отключено'}")
    doc.add_paragraph(f"Кэширование браузера: {'настроено' if h['caching_enabled'] else 'не настроено'}")
    doc.add_paragraph(f"Использование CDN: {'да' if h['cdn_used'] else 'нет'}")
    doc.add_paragraph(f"HTTPS: {'да' if h['https_used'] else 'нет'}")
    doc.add_paragraph(f"Количество редиректов: {h['redirects_count']}")
    
    # Минификация
    m = report_dict['optimization']['minification']
    doc.add_paragraph(f"Минификация HTML: {'да' if m['html'] else 'нет (или не определена)'}")
    doc.add_paragraph(f"Минификация CSS: {'да' if m['css'] else 'нет (рекомендуется)'}")
    doc.add_paragraph(f"Минификация JS: {'да' if m['js'] else 'нет (рекомендуется)'}")
    
    # Рекомендации
    doc.add_heading("4. Рекомендации по ускорению", level=1)
    for i, rec in enumerate(report_dict['recommendations'], 1):
        doc.add_paragraph(f"{i}. {rec}")
    
    doc.save(filename)
    print(f"DOCX отчёт сохранён в {filename}")

# ----------------------------------------------------------------------
# Запуск
# ----------------------------------------------------------------------
if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Анализ скорости загрузки сайта")
    parser.add_argument("url", help="URL сайта для анализа")
    parser.add_argument("--output", "-o", default="speed_audit", help="Базовое имя выходных файлов (без расширения)")
    args = parser.parse_args()
    
    target_url = args.url
    if not target_url.startswith(('http://', 'https://')):
        target_url = 'https://' + target_url
    
    result = analyze_performance(target_url)
    
    # Вывод кратких результатов в консоль
    if "error" in result:
        print(f"Ошибка: {result['error']}")
    else:
        print(f"\nВремя загрузки: {result['load_time']['full_load_sec']} с")
        print(f"TTFB: {result['load_time']['ttfb_sec']} с")
        print(f"HTTP запросов: {result['resources']['total_requests']}")
        print(f"Общий вес: {result['resources']['total_size_mb']} МБ")
        print(f"GZIP: {'вкл' if result['headers']['gzip_enabled'] else 'выкл'}")
        print(f"\nРекомендаций: {len(result['recommendations'])}")
    
    # Сохранение
    json_file = f"{args.output}.json"
    save_json(result, json_file)
    
    docx_file = f"{args.output}.docx"
    save_docx(result, docx_file)
    
    print("\nГотово.")