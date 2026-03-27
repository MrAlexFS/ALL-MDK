#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Практическая работа №4: Техническая оптимизация сайта.
Проверяет robots.txt, доступность CSS/JS для Googlebot, скорость,
мобильную оптимизацию, валидность HTML, битые ссылки, дубли.
Сохраняет отчёт в DOCX и JSON.
Использование: python tech_audit.py <URL> [--output OUTPUT_NAME]
"""

import sys
import requests
from urllib.parse import urlparse, urljoin
from bs4 import BeautifulSoup
import time
import json
import argparse
import re
from datetime import datetime
from concurrent.futures import ThreadPoolExecutor, as_completed

# Настройки
TIMEOUT = 10
USER_AGENT_DESKTOP = "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
USER_AGENT_MOBILE = "Mozilla/5.0 (iPhone; CPU iPhone OS 15_0 like Mac OS X) AppleWebKit/605.1.15 (KHTML, like Gecko) Version/15.0 Mobile/15E148 Safari/604.1"
USER_AGENT_GOOGLEBOT = "Mozilla/5.0 (compatible; Googlebot/2.1; +http://www.google.com/bot.html)"
CHECK_LINKS_LIMIT = 50

if sys.platform == 'win32':
    sys.stdout.reconfigure(encoding='utf-8')


# ----------------------------------------------------------------------
# Вспомогательные функции
# ----------------------------------------------------------------------
def get_html(url, user_agent=USER_AGENT_DESKTOP, timeout=TIMEOUT):
    """Получает HTML и заголовки ответа"""
    try:
        headers = {'User-Agent': user_agent}
        start = time.time()
        response = requests.get(url, headers=headers, timeout=timeout, allow_redirects=True)
        load_time = time.time() - start
        if response.encoding is None or response.encoding.lower() == 'iso-8859-1':
            response.encoding = response.apparent_encoding or 'utf-8'
        return response.text, response, load_time
    except Exception as e:
        return None, None, None


def check_robots_txt(base_url):
    """Проверяет robots.txt: наличие, запрещённые директории"""
    robots_url = urljoin(base_url, "/robots.txt")
    html, response, _ = get_html(robots_url)
    if not html:
        return {"exists": False, "content": None, "disallowed": []}
    
    # Парсим robots.txt (очень упрощённо)
    disallowed = []
    for line in html.splitlines():
        line = line.strip()
        if line.lower().startswith('disallow:'):
            path = line.split(':', 1)[1].strip()
            if path and path != '/':
                disallowed.append(path)
    return {
        "exists": True,
        "content": html[:500],
        "disallowed": disallowed[:10]
    }


def check_resources_accessibility(base_url):
    """Проверяет, может ли Googlebot получить доступ к CSS и JS"""
    # Загружаем главную страницу как Googlebot
    html, _, _ = get_html(base_url, USER_AGENT_GOOGLEBOT)
    if not html:
        return {"error": "Страница не загружена Googlebot"}
    
    soup = BeautifulSoup(html, 'html.parser')
    resources = []
    
    # CSS
    for link in soup.find_all('link', rel='stylesheet'):
        href = link.get('href')
        if href:
            resources.append(('css', urljoin(base_url, href)))
    
    # JS
    for script in soup.find_all('script', src=True):
        src = script['src']
        if src:
            resources.append(('js', urljoin(base_url, src)))
    
    # Проверяем доступность каждого ресурса
    blocked = []
    accessible = []
    
    for res_type, url in resources[:30]:  # ограничим для скорости
        resp = requests.head(url, headers={'User-Agent': USER_AGENT_GOOGLEBOT}, timeout=5)
        if resp.status_code == 200:
            accessible.append((res_type, url))
        else:
            blocked.append((res_type, url, resp.status_code))
    
    return {
        "total_css_js": len(resources),
        "blocked_count": len(blocked),
        "blocked_examples": blocked[:5],
        "accessible_examples": accessible[:5]
    }


def measure_speed(url):
    """Измеряет скорость загрузки (десктоп и мобильный)"""
    result = {"desktop": {}, "mobile": {}}
    
    # Десктоп
    html, resp, load_time = get_html(url, USER_AGENT_DESKTOP)
    if html:
        result["desktop"] = {
            "load_time_sec": round(load_time, 2),
            "ttfb_sec": round(resp.elapsed.total_seconds(), 2),
            "size_bytes": len(resp.content),
            "size_mb": round(len(resp.content) / (1024*1024), 2),
            "status_code": resp.status_code
        }
    else:
        result["desktop"] = {"error": "Не удалось загрузить"}
    
    # Мобильный
    html, resp, load_time = get_html(url, USER_AGENT_MOBILE)
    if html:
        result["mobile"] = {
            "load_time_sec": round(load_time, 2),
            "ttfb_sec": round(resp.elapsed.total_seconds(), 2),
            "size_bytes": len(resp.content),
            "size_mb": round(len(resp.content) / (1024*1024), 2),
            "status_code": resp.status_code
        }
    else:
        result["mobile"] = {"error": "Не удалось загрузить"}
    
    return result


def check_mobile_friendliness(url):
    """Базовая проверка адаптивности (наличие viewport, медиа-запросов)"""
    html, _, _ = get_html(url, USER_AGENT_DESKTOP)
    if not html:
        return {"error": "Не удалось загрузить страницу"}
    
    soup = BeautifulSoup(html, 'html.parser')
    viewport = soup.find('meta', attrs={'name': 'viewport'})
    has_viewport = viewport is not None
    
    # Проверяем наличие медиа-запросов в CSS (очень грубо)
    media_queries_found = False
    for link in soup.find_all('link', rel='stylesheet'):
        href = link.get('href')
        if href:
            css_url = urljoin(url, href)
            try:
                css_resp = requests.get(css_url, timeout=5)
                if '@media' in css_resp.text:
                    media_queries_found = True
                    break
            except:
                pass
    
    return {
        "viewport_meta": has_viewport,
        "media_queries_detected": media_queries_found,
        "recommendation": "OK" if has_viewport else "Добавьте <meta name='viewport'> для мобильной адаптации"
    }


def validate_html(url):
    """Базовая валидация HTML (проверка закрывающих тегов, дублирующихся ID)"""
    html, _, _ = get_html(url)
    if not html:
        return {"error": "Не удалось загрузить"}
    
    soup = BeautifulSoup(html, 'html.parser')
    errors = []
    
    # Проверка на незакрытые теги (упрощённо)
    # Простейший способ: проверить, что количество открывающих и закрывающих тегов совпадает для основных
    tags = ['div', 'p', 'h1', 'h2', 'h3', 'span']
    for tag in tags:
        open_count = len(soup.find_all(tag))
        close_count = len(re.findall(f'</{tag}>', html, re.IGNORECASE))
        if open_count != close_count:
            errors.append(f"Тег <{tag}>: открывающих {open_count}, закрывающих {close_count}")
    
    # Проверка на дублирующиеся ID
    ids = {}
    for elem in soup.find_all(id=True):
        elem_id = elem['id']
        if elem_id in ids:
            errors.append(f"Дублирующийся ID: {elem_id}")
        else:
            ids[elem_id] = True
    
    return {
        "valid": len(errors) == 0,
        "errors": errors[:10]
    }


def check_broken_links(url, limit=CHECK_LINKS_LIMIT):
    """Проверяет битые ссылки на странице"""
    html, _, _ = get_html(url)
    if not html:
        return {"error": "Не удалось загрузить страницу", "broken": []}
    
    soup = BeautifulSoup(html, 'html.parser')
    links = []
    for a in soup.find_all('a', href=True):
        href = a['href'].strip()
        if href.startswith('#') or href.startswith('javascript:') or href.startswith('mailto:'):
            continue
        full_url = urljoin(url, href)
        links.append(full_url)
    
    unique_links = list(dict.fromkeys(links))[:limit]
    broken = []
    
    # Используем ThreadPoolExecutor для ускорения
    def check_link(link):
        try:
            resp = requests.head(link, timeout=5, allow_redirects=True)
            if resp.status_code >= 400:
                return (link, resp.status_code)
        except:
            return (link, "ошибка соединения")
        return None
    
    with ThreadPoolExecutor(max_workers=10) as executor:
        futures = {executor.submit(check_link, link): link for link in unique_links}
        for future in as_completed(futures):
            result = future.result()
            if result:
                broken.append(result)
    
    return {
        "total_checked": len(unique_links),
        "broken_count": len(broken),
        "broken_examples": broken[:10]
    }


def check_duplicate_content(url):
    """Проверяет наличие дублей: www, слеш, http/https"""
    parsed = urlparse(url)
    variants = []
    # www/non-www
    if parsed.netloc.startswith('www.'):
        no_www = url.replace('www.', '', 1)
        variants.append(no_www)
    else:
        www_url = url.replace(parsed.netloc, 'www.' + parsed.netloc, 1)
        variants.append(www_url)
    # trailing slash
    if parsed.path.endswith('/'):
        no_slash = url.rstrip('/')
        variants.append(no_slash)
    else:
        slash_url = url + '/'
        variants.append(slash_url)
    # http/https
    if parsed.scheme == 'https':
        http_url = 'http://' + parsed.netloc + parsed.path
        variants.append(http_url)
    else:
        https_url = 'https://' + parsed.netloc + parsed.path
        variants.append(https_url)
    
    accessible = []
    for v in set(variants):
        if v == url:
            continue
        r, _, _ = get_html(v)
        if r:
            accessible.append(v)
    
    return {
        "potential_duplicates": list(set(variants)),
        "accessible_duplicates": accessible
    }


def generate_recommendations(data):
    """Генерирует рекомендации на основе собранных данных"""
    recs = []
    
    # robots.txt
    if not data['robots']['exists']:
        recs.append("Создайте robots.txt и разместите в корне сайта.")
    elif data['robots']['disallowed']:
        recs.append("Проверьте, не заблокированы ли важные разделы в robots.txt.")
    
    # Доступность CSS/JS для Googlebot
    res = data['resources_accessibility']
    if res.get('blocked_count', 0) > 0:
        recs.append(f"Googlebot не может получить доступ к {res['blocked_count']} ресурсам CSS/JS. Разрешите индексацию в robots.txt.")
    
    # Скорость
    speed = data['speed']
    if 'desktop' in speed and 'load_time_sec' in speed['desktop']:
        if speed['desktop']['load_time_sec'] > 3:
            recs.append(f"Время загрузки на десктопе: {speed['desktop']['load_time_sec']} с. Оптимизируйте изображения, включите сжатие, кэширование.")
        if speed['desktop'].get('size_mb', 0) > 2:
            recs.append(f"Вес страницы {speed['desktop']['size_mb']} МБ. Сжимайте изображения, минифицируйте код.")
    if 'mobile' in speed and 'load_time_sec' in speed['mobile'] and speed['mobile']['load_time_sec'] > 3:
        recs.append(f"Время загрузки на мобильных: {speed['mobile']['load_time_sec']} с. Ускорьте мобильную версию.")
    
    # Мобильная оптимизация
    mobile = data['mobile_friendly']
    if mobile.get('error'):
        recs.append("Не удалось проверить мобильную оптимизацию.")
    elif not mobile.get('viewport_meta'):
        recs.append("Добавьте мета-тег viewport для корректного отображения на мобильных устройствах.")
    if not mobile.get('media_queries_detected'):
        recs.append("Используйте CSS-медиазапросы для адаптивной вёрстки.")
    
    # Валидность HTML
    html_valid = data['html_validation']
    if not html_valid.get('valid'):
        recs.append(f"Найдены ошибки валидации HTML: {len(html_valid.get('errors', []))}. Исправьте их.")
    
    # Битые ссылки
    broken = data['broken_links']
    if broken.get('broken_count', 0) > 0:
        recs.append(f"Обнаружено {broken['broken_count']} битых ссылок. Исправьте или удалите их.")
    
    # Дубли
    dup = data['duplicates']
    if dup.get('accessible_duplicates'):
        recs.append(f"Найдены доступные зеркала: {', '.join(dup['accessible_duplicates'])}. Настройте 301-редирект.")
    
    return recs


def technical_audit(url):
    """Основная функция технического аудита"""
    print(f"Провожу технический аудит: {url}")
    report = {
        "url": url,
        "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "robots": {},
        "resources_accessibility": {},
        "speed": {},
        "mobile_friendly": {},
        "html_validation": {},
        "broken_links": {},
        "duplicates": {},
        "recommendations": []
    }
    
    # 1. robots.txt
    report["robots"] = check_robots_txt(url)
    
    # 2. Доступность CSS/JS для Googlebot
    report["resources_accessibility"] = check_resources_accessibility(url)
    
    # 3. Скорость загрузки (десктоп + мобильный)
    report["speed"] = measure_speed(url)
    
    # 4. Мобильная оптимизация
    report["mobile_friendly"] = check_mobile_friendliness(url)
    
    # 5. Валидность HTML
    report["html_validation"] = validate_html(url)
    
    # 6. Битые ссылки
    report["broken_links"] = check_broken_links(url)
    
    # 7. Дубли контента (зеркала)
    report["duplicates"] = check_duplicate_content(url)
    
    # 8. Рекомендации
    report["recommendations"] = generate_recommendations(report)
    
    return report


def save_json(data, filename):
    with open(filename, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    print(f"JSON отчёт сохранён в {filename}")


def save_docx(data, filename):
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("Для сохранения DOCX установите python-docx: pip install python-docx")
        return
    
    doc = Document()
    heading = doc.add_heading("Отчёт по технической оптимизации сайта", 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    doc.add_paragraph(f"URL: {data['url']}")
    doc.add_paragraph(f"Дата анализа: {data['timestamp']}")
    doc.add_paragraph()
    
    # 1. Robots.txt
    doc.add_heading("1. Robots.txt", level=1)
    r = data['robots']
    doc.add_paragraph(f"Наличие: {'Да' if r['exists'] else 'Нет'}")
    if r['exists']:
        doc.add_paragraph(f"Содержимое (первые 500 символов):\n{r['content']}")
        if r['disallowed']:
            doc.add_paragraph(f"Запрещённые директории: {', '.join(r['disallowed'])}")
    
    # 2. Доступность CSS/JS для Googlebot
    doc.add_heading("2. Доступность CSS и JS для Googlebot", level=1)
    res = data['resources_accessibility']
    if 'error' in res:
        doc.add_paragraph(f"Ошибка: {res['error']}")
    else:
        doc.add_paragraph(f"Всего CSS/JS ресурсов: {res['total_css_js']}")
        doc.add_paragraph(f"Заблокировано для Googlebot: {res['blocked_count']}")
        if res['blocked_examples']:
            doc.add_paragraph("Примеры заблокированных ресурсов:")
            for typ, url, code in res['blocked_examples']:
                doc.add_paragraph(f"  {typ}: {url} (код {code})")
    
    # 3. Скорость загрузки
    doc.add_heading("3. Скорость загрузки", level=1)
    speed = data['speed']
    for device in ['desktop', 'mobile']:
        doc.add_heading(f"{device.upper()}", level=2)
        s = speed.get(device, {})
        if 'error' in s:
            doc.add_paragraph(f"Ошибка: {s['error']}")
        else:
            doc.add_paragraph(f"Время загрузки: {s['load_time_sec']} с")
            doc.add_paragraph(f"TTFB: {s['ttfb_sec']} с")
            doc.add_paragraph(f"Вес страницы: {s['size_mb']} МБ")
            doc.add_paragraph(f"HTTP статус: {s['status_code']}")
    
    # 4. Мобильная оптимизация
    doc.add_heading("4. Мобильная оптимизация", level=1)
    mob = data['mobile_friendly']
    if 'error' in mob:
        doc.add_paragraph(f"Ошибка: {mob['error']}")
    else:
        doc.add_paragraph(f"Мета-тег viewport: {'есть' if mob['viewport_meta'] else 'нет'}")
        doc.add_paragraph(f"Медиазапросы в CSS: {'обнаружены' if mob['media_queries_detected'] else 'не обнаружены'}")
        doc.add_paragraph(f"Рекомендация: {mob['recommendation']}")
    
    # 5. Валидность HTML
    doc.add_heading("5. Валидность HTML", level=1)
    html_val = data['html_validation']
    if 'error' in html_val:
        doc.add_paragraph(f"Ошибка: {html_val['error']}")
    else:
        doc.add_paragraph(f"Статус: {'Валидно' if html_val['valid'] else 'Найдены ошибки'}")
        if html_val['errors']:
            doc.add_paragraph("Ошибки:")
            for err in html_val['errors'][:10]:
                doc.add_paragraph(f"  - {err}")
    
    # 6. Битые ссылки
    doc.add_heading("6. Битые ссылки", level=1)
    broken = data['broken_links']
    if 'error' in broken:
        doc.add_paragraph(f"Ошибка: {broken['error']}")
    else:
        doc.add_paragraph(f"Проверено ссылок: {broken['total_checked']}")
        doc.add_paragraph(f"Битых ссылок: {broken['broken_count']}")
        if broken['broken_examples']:
            doc.add_paragraph("Примеры битых ссылок:")
            for link, code in broken['broken_examples'][:5]:
                doc.add_paragraph(f"  {link} -> {code}")
    
    # 7. Дубли контента
    doc.add_heading("7. Дубли контента (зеркала)", level=1)
    dup = data['duplicates']
    if dup['potential_duplicates']:
        doc.add_paragraph("Потенциальные дубли (зеркала):")
        for v in dup['potential_duplicates']:
            doc.add_paragraph(f"  {v}")
    if dup['accessible_duplicates']:
        doc.add_paragraph("Доступные зеркала (требуется 301-редирект):")
        for a in dup['accessible_duplicates']:
            doc.add_paragraph(f"  {a}")
    else:
        doc.add_paragraph("Активных зеркал не обнаружено.")
    
    # 8. Рекомендации
    doc.add_heading("8. Рекомендации по технической оптимизации", level=1)
    for i, rec in enumerate(data['recommendations'], 1):
        doc.add_paragraph(f"{i}. {rec}")
    
    doc.save(filename)
    print(f"DOCX отчёт сохранён в {filename}")


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Технический аудит сайта")
    parser.add_argument("url", help="URL сайта для анализа")
    parser.add_argument("--output", "-o", default="tech_audit", help="Базовое имя выходных файлов")
    args = parser.parse_args()
    
    target_url = args.url
    if not target_url.startswith(('http://', 'https://')):
        target_url = 'https://' + target_url
    
    result = technical_audit(target_url)
    
    if "error" in result:
        print(f"Ошибка: {result['error']}")
    else:
        print(f"\nАнализ завершён. Найдено {len(result['recommendations'])} рекомендаций.")
    
    save_json(result, f"{args.output}.json")
    save_docx(result, f"{args.output}.docx")
    print("\nГотово.")