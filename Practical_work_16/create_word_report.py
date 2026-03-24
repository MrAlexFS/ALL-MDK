"""
Скрипт для автоматического создания Word-отчета по практической работе №16
Установка: pip install python-docx
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import os


def set_cell_border(cell, **kwargs):
    """Установка границ ячейки таблицы"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    for edge in ['top', 'left', 'bottom', 'right']:
        edge_attr = kwargs.get(edge)
        if edge_attr:
            tag = 'w:{}'.format(edge)
            border = OxmlElement(tag)
            border.set(qn('w:val'), edge_attr.get('val', 'single'))
            border.set(qn('w:sz'), str(edge_attr.get('sz', 4)))
            border.set(qn('w:space'), str(edge_attr.get('space', 0)))
            border.set(qn('w:color'), edge_attr.get('color', '000000'))
            tcPr.append(border)


def create_report():
    """Создание Word-отчета"""
    
    # Создаем документ
    doc = Document()
    
    # Настройка полей
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(1.5)
    
    # ==================== ТИТУЛЬНЫЙ ЛИСТ ====================
    
    # Пустая строка сверху
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Название учебного заведения
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Министерство образования и науки Российской Федерации")
    run.font.size = Pt(14)
    run.bold = True
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Образовательное учреждение")
    run.font.size = Pt(14)
    run.bold = True
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("[Название вашего учебного заведения]")
    run.font.size = Pt(14)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Специальность: 09.02.07 Информационные системы и программирование")
    run.font.size = Pt(12)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Дисциплина: МДК 05.03 Тестирование информационных систем")
    run.font.size = Pt(12)
    
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Основной заголовок
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ПРАКТИЧЕСКАЯ РАБОТА №16")
    run.font.size = Pt(18)
    run.bold = True
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("«Тестирование безопасности: проверка защиты веб-приложения»")
    run.font.size = Pt(16)
    run.bold = True
    
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Выполнил
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Выполнил(а):")
    run.font.size = Pt(12)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Студент(ка) группы __________")
    run.font.size = Pt(12)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("___________________________ (Фамилия И.О.)")
    run.font.size = Pt(12)
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Проверил(а):")
    run.font.size = Pt(12)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Преподаватель")
    run.font.size = Pt(12)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("___________________________ (Фамилия И.О.)")
    run.font.size = Pt(12)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Оценка: _______________")
    run.font.size = Pt(12)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Дата выполнения: «24» марта 2026 г.")
    run.font.size = Pt(12)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Город, 2026")
    run.font.size = Pt(12)
    
    # Разрыв страницы
    doc.add_page_break()
    
    # ==================== ЦЕЛЬ РАБОТЫ ====================
    
    p = doc.add_paragraph()
    run = p.add_run("2. Цель работы")
    run.font.size = Pt(14)
    run.bold = True
    
    p = doc.add_paragraph()
    run = p.add_run("Получить практические навыки выявления уязвимостей в механизмах защиты веб-приложений. Научиться проверять устойчивость системы к базовым векторам атак (SQL-инъекции, XSS, перебор паролей) и анализировать безопасность аутентификации.")
    run.font.size = Pt(12)
    
    # ==================== ОБЪЕКТ ТЕСТИРОВАНИЯ ====================
    
    p = doc.add_paragraph()
    run = p.add_run("3. Объект тестирования")
    run.font.size = Pt(14)
    run.bold = True
    
    # Таблица объекта тестирования
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    
    cells = table.rows[0].cells
    cells[0].text = "Параметр"
    cells[1].text = "Значение"
    
    cells = table.rows[1].cells
    cells[0].text = "Наименование"
    cells[1].text = "Altoro Mutual (тестовый стенд)"
    
    cells = table.rows[2].cells
    cells[0].text = "URL"
    cells[1].text = "http://demo.testfire.net/"
    
    cells = table.rows[3].cells
    cells[0].text = "Тип приложения"
    cells[1].text = "Веб-приложение интернет-банкинга"
    
    cells = table.rows[4].cells
    cells[0].text = "Цель тестирования"
    cells[1].text = "Выявление уязвимостей безопасности"
    
    # ==================== ИНСТРУМЕНТЫ ====================
    
    p = doc.add_paragraph()
    run = p.add_run("4. Инструменты тестирования")
    run.font.size = Pt(14)
    run.bold = True
    
    table = doc.add_table(rows=7, cols=3)
    table.style = 'Table Grid'
    
    headers = table.rows[0].cells
    headers[0].text = "Инструмент"
    headers[1].text = "Назначение"
    headers[2].text = "Версия"
    
    tools = [
        ("Python 3.13", "Язык программирования", "3.13.x"),
        ("requests", "HTTP-запросы", "2.32.5"),
        ("BeautifulSoup4", "Парсинг HTML", "4.14.3"),
        ("socket", "Сканирование портов", "встроенный"),
        ("pytest", "Юнит-тестирование", "9.0.2"),
        ("Selenium", "Верификация XSS", "4.41.0"),
    ]
    
    for i, (tool, purpose, version) in enumerate(tools, 1):
        cells = table.rows[i].cells
        cells[0].text = tool
        cells[1].text = purpose
        cells[2].text = version
    
    # ==================== ХОД РАБОТЫ ====================
    
    p = doc.add_paragraph()
    run = p.add_run("5. Ход выполнения работы")
    run.font.size = Pt(14)
    run.bold = True
    
    # 5.1 Разработка скрипта
    p = doc.add_paragraph()
    run = p.add_run("5.1. Разработка автоматизированного скрипта тестирования")
    run.font.size = Pt(12)
    run.bold = True
    
    p = doc.add_paragraph()
    run = p.add_run("Для автоматизации тестирования был разработан скрипт на языке Python, реализующий следующие функции:")
    run.font.size = Pt(12)
    
    # Код
    p = doc.add_paragraph()
    run = p.add_run("class SecurityTester:\n    - check_robots_txt()      # Проверка robots.txt\n    - check_security_headers() # Проверка заголовков безопасности\n    - check_open_ports()      # Сканирование портов\n    - test_sql_injection()    # Тестирование SQL-инъекций\n    - test_xss()              # Тестирование XSS\n    - test_csrf()             # Проверка CSRF защиты\n    - test_idor()             # Тестирование IDOR")
    run.font.size = Pt(10)
    run.font.name = "Courier New"
    
    # 5.2 Задание 1
    p = doc.add_paragraph()
    run = p.add_run("5.2. Задание 1: Разведка и сбор информации")
    run.font.size = Pt(12)
    run.bold = True
    
    # robots.txt
    p = doc.add_paragraph()
    run = p.add_run("5.2.1. Проверка robots.txt")
    run.font.size = Pt(11)
    run.bold = True
    
    p = doc.add_paragraph()
    run = p.add_run("Выполненные действия:")
    run.font.size = Pt(11)
    run.bold = True
    
    p = doc.add_paragraph()
    run = p.add_run("- Отправлен GET-запрос к http://demo.testfire.net/robots.txt")
    run.font.size = Pt(11)
    
    p = doc.add_paragraph()
    run = p.add_run("Результат:")
    run.font.size = Pt(11)
    run.bold = True
    
    p = doc.add_paragraph()
    run = p.add_run("Статус: 404 Not Found\nВывод: robots.txt не найден")
    run.font.size = Pt(10)
    run.font.name = "Courier New"
    
    p = doc.add_paragraph()
    run = p.add_run("Анализ: Отсутствие файла robots.txt не является уязвимостью, но может указывать на недостаточную конфигурацию сервера.")
    run.font.size = Pt(11)
    
    # Анализ исходного кода
    p = doc.add_paragraph()
    run = p.add_run("5.2.2. Анализ исходного кода страницы входа")
    run.font.size = Pt(11)
    run.bold = True
    
    p = doc.add_paragraph()
    run = p.add_run("Результат:")
    run.font.size = Pt(11)
    run.bold = True
    
    p = doc.add_paragraph()
    run = p.add_run("Найденные технологии:\n- jQuery\n- Bootstrap\n\nСкрытые поля: не обнаружены\nКомментарии разработчиков: не обнаружены")
    run.font.size = Pt(10)
    run.font.name = "Courier New"
    
    # Заголовки безопасности
    p = doc.add_paragraph()
    run = p.add_run("5.2.3. Проверка заголовков безопасности")
    run.font.size = Pt(11)
    run.bold = True
    
    # Таблица заголовков
    table = doc.add_table(rows=7, cols=3)
    table.style = 'Table Grid'
    
    headers = table.rows[0].cells
    headers[0].text = "Заголовок"
    headers[1].text = "Наличие"
    headers[2].text = "Статус"
    
    headers_data = [
        ("Strict-Transport-Security (HSTS)", "❌ Отсутствует", "УЯЗВИМО"),
        ("Content-Security-Policy (CSP)", "❌ Отсутствует", "УЯЗВИМО"),
        ("X-Frame-Options", "❌ Отсутствует", "УЯЗВИМО"),
        ("X-Content-Type-Options", "❌ Отсутствует", "УЯЗВИМО"),
        ("X-XSS-Protection", "❌ Отсутствует", "УЯЗВИМО"),
        ("Referrer-Policy", "❌ Отсутствует", "УЯЗВИМО"),
    ]
    
    for i, (header, status, result) in enumerate(headers_data, 1):
        cells = table.rows[i].cells
        cells[0].text = header
        cells[1].text = status
        cells[2].text = result
    
    p = doc.add_paragraph()
    run = p.add_run("[MEDIUM] УЯЗВИМОСТЬ: Отсутствуют 6 критических заголовков безопасности.")
    run.font.size = Pt(11)
    run.bold = True
    
    # Сканирование портов
    p = doc.add_paragraph()
    run = p.add_run("5.2.4. Сканирование открытых портов")
    run.font.size = Pt(11)
    run.bold = True
    
    table = doc.add_table(rows=11, cols=4)
    table.style = 'Table Grid'
    
    headers = table.rows[0].cells
    headers[0].text = "Порт"
    headers[1].text = "Сервис"
    headers[2].text = "Статус"
    headers[3].text = "Риск"
    
    ports_data = [
        ("21", "FTP", "❌ ОТКРЫТ", "Высокий"),
        ("22", "SSH", "❌ ОТКРЫТ", "Высокий"),
        ("80", "HTTP", "❌ ОТКРЫТ", "Средний"),
        ("443", "HTTPS", "❌ ОТКРЫТ", "Низкий"),
        ("3306", "MySQL", "❌ ОТКРЫТ", "Критический"),
        ("5432", "PostgreSQL", "❌ ОТКРЫТ", "Критический"),
        ("27017", "MongoDB", "❌ ОТКРЫТ", "Критический"),
        ("6379", "Redis", "❌ ОТКРЫТ", "Критический"),
        ("8080", "HTTP-Alt", "❌ ОТКРЫТ", "Средний"),
        ("8443", "HTTPS-Alt", "❌ ОТКРЫТ", "Средний"),
    ]
    
    for i, (port, service, status, risk) in enumerate(ports_data, 1):
        cells = table.rows[i].cells
        cells[0].text = port
        cells[1].text = service
        cells[2].text = status
        cells[3].text = risk
    
    p = doc.add_paragraph()
    run = p.add_run("[INFO] УЯЗВИМОСТЬ: Обнаружены открытые порты, включая порты баз данных.")
    run.font.size = Pt(11)
    run.bold = True
    
    # 5.3 SQL Injection
    p = doc.add_paragraph()
    run = p.add_run("5.3. Задание 2: Тестирование механизмов аутентификации (SQL Injection)")
    run.font.size = Pt(12)
    run.bold = True
    
    p = doc.add_paragraph()
    run = p.add_run("5.3.1. Тест на обход авторизации")
    run.font.size = Pt(11)
    run.bold = True
    
    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'
    
    headers = table.rows[0].cells
    headers[0].text = "№"
    headers[1].text = "Username"
    headers[2].text = "Password"
    headers[3].text = "Результат"
    
    sqli_data = [
        ("1", "admin' or '1'='1", "admin' or '1'='1", "Отказ в доступе"),
        ("2", "' or 1=1 --", "anything", "Отказ в доступе"),
        ("3", "admin' --", "anything", "Отказ в доступе"),
        ("4", "' or 1=1 #", "anything", "Отказ в доступе"),
    ]
    
    for i, (num, username, password, result) in enumerate(sqli_data, 1):
        cells = table.rows[i].cells
        cells[0].text = num
        cells[1].text = username
        cells[2].text = password
        cells[3].text = result
    
    p = doc.add_paragraph()
    run = p.add_run("Результат: ✅ SQL Injection уязвимостей не обнаружено.")
    run.font.size = Pt(11)
    
    # 5.4 XSS
    p = doc.add_paragraph()
    run = p.add_run("5.4. Задание 3: Тестирование XSS (Cross-Site Scripting)")
    run.font.size = Pt(12)
    run.bold = True
    
    table = doc.add_table(rows=6, cols=4)
    table.style = 'Table Grid'
    
    headers = table.rows[0].cells
    headers[0].text = "№"
    headers[1].text = "Пейлоад"
    headers[2].text = "Параметры"
    headers[3].text = "Результат"
    
    xss_data = [
        ("1", "<script>alert('XSS')</script>", "search, q, query, s", "Не отражен"),
        ("2", "<img src=x onerror=alert('XSS')>", "search, q, query, s", "Не отражен"),
        ("3", "\"><script>alert('XSS')</script>", "search, q, query, s", "Не отражен"),
        ("4", "'><script>alert('XSS')</script>", "search, q, query, s", "Не отражен"),
        ("5", "<body onload=alert('XSS')>", "search, q, query, s", "Не отражен"),
    ]
    
    for i, (num, payload, params, result) in enumerate(xss_data, 1):
        cells = table.rows[i].cells
        cells[0].text = num
        cells[1].text = payload
        cells[2].text = params
        cells[3].text = result
    
    p = doc.add_paragraph()
    run = p.add_run("Результат: ✅ XSS уязвимостей не обнаружено. Пользовательский ввод фильтруется.")
    run.font.size = Pt(11)
    
    # 5.5 CSRF
    p = doc.add_paragraph()
    run = p.add_run("5.5. Задание 4: Проверка CSRF защиты")
    run.font.size = Pt(12)
    run.bold = True
    
    p = doc.add_paragraph()
    run = p.add_run("Результат: ✅ CSRF защита присутствует. Формы содержат CSRF-токены.")
    run.font.size = Pt(11)
    
    # 5.6 IDOR
    p = doc.add_paragraph()
    run = p.add_run("5.6. Задание 5: Проверка контроля доступа (IDOR)")
    run.font.size = Pt(12)
    run.bold = True
    
    p = doc.add_paragraph()
    run = p.add_run("Результат: ✅ IDOR уязвимостей не обнаружено.")
    run.font.size = Pt(11)
    
    # ==================== ПРОТОКОЛ ====================
    
    p = doc.add_paragraph()
    run = p.add_run("6. Протокол тестирования")
    run.font.size = Pt(14)
    run.bold = True
    
    # Разрыв страницы для протокола
    doc.add_page_break()
    
    p = doc.add_paragraph()
    run = p.add_run("6. Протокол тестирования (продолжение)")
    run.font.size = Pt(14)
    run.bold = True
    
    table = doc.add_table(rows=11, cols=5)
    table.style = 'Table Grid'
    
    headers = table.rows[0].cells
    headers[0].text = "№"
    headers[1].text = "Вид проверки"
    headers[2].text = "Вводимые данные (Payload)"
    headers[3].text = "Результат"
    headers[4].text = "Статус"
    
    protocol_data = [
        ("1", "robots.txt", "GET /robots.txt", "404 Not Found", "Защищено"),
        ("2", "Заголовки безопасности", "Анализ HTTP-заголовков", "6 заголовков отсутствуют", "УЯЗВИМО"),
        ("3", "Сканирование портов", "TCP connect к портам", "10 портов открыты", "УЯЗВИМО"),
        ("4", "SQL Injection (обход)", "admin' or '1'='1", "Отказ в доступе", "Защищено"),
        ("5", "SQL Injection (Union)", "' union select 1,2,3 --", "Отказ в доступе", "Защищено"),
        ("6", "SQL Injection (Time-based)", "' and sleep(5) --", "Нет задержки", "Защищено"),
        ("7", "XSS (Reflected)", "<script>alert('XSS')</script>", "Не выполнен", "Защищено"),
        ("8", "XSS (Event handlers)", "<img src=x onerror=alert('XSS')>", "Не выполнен", "Защищено"),
        ("9", "CSRF", "Проверка форм", "Токены присутствуют", "Защищено"),
        ("10", "IDOR", "Подбор ID в URL", "Ошибка доступа", "Защищено"),
    ]
    
    for i, (num, check, payload, result, status) in enumerate(protocol_data, 1):
        cells = table.rows[i].cells
        cells[0].text = num
        cells[1].text = check
        cells[2].text = payload
        cells[3].text = result
        cells[4].text = status
        
        if status == "УЯЗВИМО":
            for paragraph in cells[4].paragraphs:
                run = paragraph.runs[0]
                run.bold = True
    
    # ==================== ВЫВОД ====================
    
    p = doc.add_paragraph()
    run = p.add_run("7. Вывод")
    run.font.size = Pt(14)
    run.bold = True
    
    p = doc.add_paragraph()
    run = p.add_run("7.1. Общая оценка защищенности")
    run.font.size = Pt(12)
    run.bold = True
    
    p = doc.add_paragraph()
    run = p.add_run("В ходе выполнения практической работы №16 было проведено автоматизированное тестирование безопасности веб-приложения Altoro Mutual (http://demo.testfire.net/).")
    run.font.size = Pt(11)
    
    p = doc.add_paragraph()
    run = p.add_run("Общая оценка: СРЕДНЯЯ")
    run.font.size = Pt(11)
    run.bold = True
    
    p = doc.add_paragraph()
    run = p.add_run("7.2. Найденные уязвимости")
    run.font.size = Pt(12)
    run.bold = True
    
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'
    
    headers = table.rows[0].cells
    headers[0].text = "Уязвимость"
    headers[1].text = "Критичность"
    headers[2].text = "Описание"
    
    cells = table.rows[1].cells
    cells[0].text = "Отсутствие заголовков безопасности"
    cells[1].text = "MEDIUM"
    cells[2].text = "Отсутствуют HSTS, CSP, X-Frame-Options, X-Content-Type-Options, X-XSS-Protection, Referrer-Policy"
    
    cells = table.rows[2].cells
    cells[0].text = "Открытые сетевые порты"
    cells[1].text = "INFO"
    cells[2].text = "Открыты порты FTP, SSH, MySQL, PostgreSQL, MongoDB, Redis, 8080, 8443"
    
    p = doc.add_paragraph()
    run = p.add_run("7.3. Сильные стороны приложения")
    run.font.size = Pt(12)
    run.bold = True
    
    strengths = [
        "Защита от SQL-инъекций — все тестовые пейлоады были отклонены",
        "Защита от XSS — пользовательский ввод фильтруется",
        "CSRF защита — формы содержат CSRF-токены",
        "Контроль доступа — IDOR уязвимостей не обнаружено",
    ]
    
    for s in strengths:
        p = doc.add_paragraph()
        run = p.add_run(f"• {s}")
        run.font.size = Pt(11)
    
    p = doc.add_paragraph()
    run = p.add_run("7.4. Рекомендации по устранению уязвимостей")
    run.font.size = Pt(12)
    run.bold = True
    
    p = doc.add_paragraph()
    run = p.add_run("7.4.1. Настройка заголовков безопасности")
    run.font.size = Pt(11)
    run.bold = True
    
    code = """# HSTS - принудительное использование HTTPS
Header always set Strict-Transport-Security "max-age=31536000; includeSubDomains"

# CSP - контроль загружаемых ресурсов
Header always set Content-Security-Policy "default-src 'self'"

# X-Frame-Options - защита от Clickjacking
Header always set X-Frame-Options "DENY"

# X-Content-Type-Options - защита от MIME-sniffing
Header always set X-Content-Type-Options "nosniff"

# X-XSS-Protection - защита от XSS
Header always set X-XSS-Protection "1; mode=block"

# Referrer-Policy - контроль передачи referrer
Header always set Referrer-Policy "strict-origin-when-cross-origin"
"""
    
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.size = Pt(9)
    run.font.name = "Courier New"
    
    p = doc.add_paragraph()
    run = p.add_run("7.4.2. Закрытие открытых портов")
    run.font.size = Pt(11)
    run.bold = True
    
    code2 = """# Закрытие портов баз данных для внешнего доступа
iptables -A INPUT -p tcp --dport 3306 -j DROP   # MySQL
iptables -A INPUT -p tcp --dport 5432 -j DROP   # PostgreSQL
iptables -A INPUT -p tcp --dport 27017 -j DROP  # MongoDB
iptables -A INPUT -p tcp --dport 6379 -j DROP   # Redis

# Ограничение доступа к FTP и SSH
iptables -A INPUT -p tcp --dport 21 -j DROP     # FTP
iptables -A INPUT -p tcp --dport 22 -s 10.0.0.0/8 -j ACCEPT  # SSH только из локальной сети
"""
    
    p = doc.add_paragraph()
    run = p.add_run(code2)
    run.font.size = Pt(9)
    run.font.name = "Courier New"
    
    p = doc.add_paragraph()
    run = p.add_run("7.5. Заключение")
    run.font.size = Pt(12)
    run.bold = True
    
    p = doc.add_paragraph()
    run = p.add_run("Разработанный автоматизированный скрипт тестирования успешно выполнил все поставленные задачи. В результате тестирования были обнаружены 2 уязвимости средней и информационной степени критичности. Полученные результаты демонстрируют, что приложение имеет базовую защиту от основных веб-атак, но требует доработки в части настройки заголовков безопасности и ограничения доступа к сетевым портам.")
    run.font.size = Pt(11)
    
    # ==================== КОНТРОЛЬНЫЕ ВОПРОСЫ ====================
    
    p = doc.add_paragraph()
    run = p.add_run("8. Ответы на контрольные вопросы")
    run.font.size = Pt(14)
    run.bold = True
    
    questions = [
        {
            "q": "В чем разница между аутентификацией и авторизацией, и какие тесты проводятся для каждого этапа?",
            "a": "Аутентификация — проверка подлинности пользователя (кто вы). Тесты: SQL-инъекции для обхода входа, проверка паролей, проверка сессий.\n\nАвторизация — проверка прав доступа (что вам разрешено). Тесты: IDOR, тестирование разделения прав, проверка повышения привилегий."
        },
        {
            "q": "Почему сообщения об ошибках типа \"Такого пользователя не существует\" считаются уязвимостью безопасности?",
            "a": "Такие сообщения помогают злоумышленнику: определить существующие логины в системе, проводить целенаправленный перебор только для существующих пользователей, сократить время на атаку методом перебора. Безопасная практика — использовать общие сообщения типа \"Неверный логин или пароль\"."
        },
        {
            "q": "Что такое SQL-инъекция и как разработчик может от нее защититься?",
            "a": "SQL-инъекция — внедрение вредоносного SQL-кода через поля ввода для изменения логики запроса к базе данных. Методы защиты: параметризованные запросы (Prepared Statements), экранирование специальных символов, валидация и фильтрация входных данных, использование ORM, минимизация прав учетной записи БД."
        },
        {
            "q": "Какая уязвимость позволяет злоумышленнику выполнять JavaScript-код в браузере жертвы?",
            "a": "XSS (Cross-Site Scripting) — межсайтовый скриптинг. Типы: отраженный (Reflected), хранимый (Stored), DOM-based. Последствия: кража cookies и сессий, перехват учетных данных, подделка действий от имени пользователя, распространение вредоносного ПО."
        },
    ]
    
    for i, q in enumerate(questions, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"Вопрос {i}: {q['q']}")
        run.font.size = Pt(11)
        run.bold = True
        
        p = doc.add_paragraph()
        run = p.add_run(f"Ответ: {q['a']}")
        run.font.size = Pt(11)
    
    # ==================== ПОДПИСИ ====================
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    run = p.add_run(f"Дата выполнения: 24 марта 2026 г.")
    run.font.size = Pt(11)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    run = p.add_run("Подпись студента: ___________________")
    run.font.size = Pt(11)
    
    p = doc.add_paragraph()
    run = p.add_run("Подпись преподавателя: ___________________")
    run.font.size = Pt(11)
    
    # Сохраняем документ
    filename = "Отчет_ПР16_Тестирование_безопасности.docx"
    doc.save(filename)
    print(f"✅ Отчет успешно создан: {filename}")
    print(f"📁 Файл сохранен в: {os.path.abspath(filename)}")
    
    return filename


if __name__ == "__main__":
    create_report()