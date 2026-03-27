#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Практическая работа №3: Внутренняя SEO-оптимизация сайта.
Расширенная версия: включает проверки robots.txt, мета-тегов, дублей,
скорости, SMO, Open Graph, nofollow, кэширования, gzip.
Для пунктов, требующих ручной оценки (дизайн, уникальность текста),
добавляет в отчёт рекомендации с инструкциями.
Сохраняет отчёт в DOCX и JSON.
Использование: python seo_audit_full.py <URL> [--output OUTPUT_NAME]
"""

import sys
import requests
from urllib.parse import urlparse, urljoin
from bs4 import BeautifulSoup
import time
import json
import argparse
import re
from datetime import datetime

# Настройки
TIMEOUT = 10
USER_AGENT = "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36"
CHECK_LINKS_LIMIT = 30

# Для корректного вывода в консоль (Windows)
if sys.platform == 'win32':
    sys.stdout.reconfigure(encoding='utf-8')

# ----------------------------------------------------------------------
# Вспомогательные функции
# ----------------------------------------------------------------------
def get_html(url):
    """Получает HTML и информацию о загрузке"""
    try:
        headers = {'User-Agent': USER_AGENT}
        start = time.time()
        response = requests.get(url, headers=headers, timeout=TIMEOUT, allow_redirects=True)
        load_time = time.time() - start
        if response.encoding is None or response.encoding.lower() == 'iso-8859-1':
            response.encoding = response.apparent_encoding or 'utf-8'
        return response.text, response, load_time
    except Exception as e:
        print(f"Ошибка загрузки {url}: {e}")
        return None, None, None

def check_robots_txt(base_url):
    """Проверяет наличие и содержимое robots.txt"""
    robots_url = urljoin(base_url, "/robots.txt")
    html, _, _ = get_html(robots_url)
    return {
        "exists": html is not None,
        "content": html[:500] if html else None
    }

def analyze_meta(html):
    """Анализирует мета-теги страницы"""
    soup = BeautifulSoup(html, 'html.parser')
    result = {}
    # Title
    title_tag = soup.find('title')
    result['title'] = title_tag.get_text(strip=True) if title_tag else None
    result['title_length'] = len(result['title']) if result['title'] else 0
    # Meta description
    meta_desc = soup.find('meta', attrs={'name': 'description'})
    result['description'] = meta_desc.get('content') if meta_desc else None
    # Meta keywords
    meta_keys = soup.find('meta', attrs={'name': 'keywords'})
    result['keywords'] = meta_keys.get('content') if meta_keys else None
    # Meta robots
    meta_robots = soup.find('meta', attrs={'name': 'robots'})
    result['robots_meta'] = meta_robots.get('content') if meta_robots else None
    # Canonical
    canonical = soup.find('link', attrs={'rel': 'canonical'})
    result['canonical'] = canonical.get('href') if canonical else None
    # Open Graph (og:title, og:description, og:image, og:url)
    og = {}
    for prop in ['og:title', 'og:description', 'og:image', 'og:url']:
        tag = soup.find('meta', attrs={'property': prop})
        og[prop] = tag.get('content') if tag else None
    result['open_graph'] = og
    return result

def analyze_headings(html):
    """Анализирует заголовки H1-H6"""
    soup = BeautifulSoup(html, 'html.parser')
    headings = {}
    for i in range(1, 7):
        tag = f'h{i}'
        elements = soup.find_all(tag)
        headings[tag] = [elem.get_text(strip=True) for elem in elements]
    return headings

def analyze_images(html, base_url):
    """Анализирует изображения: количество, alt, title, форматы"""
    soup = BeautifulSoup(html, 'html.parser')
    images = soup.find_all('img')
    total = len(images)
    missing_alt = sum(1 for img in images if not img.get('alt'))
    missing_title = sum(1 for img in images if not img.get('title'))
    # Примеры названий файлов
    filenames = [img.get('src', '').split('/')[-1] for img in images[:5]]
    return {
        "total": total,
        "missing_alt": missing_alt,
        "missing_title": missing_title,
        "filenames_examples": filenames
    }

def analyze_text(html):
    """Анализирует текстовое содержимое"""
    soup = BeautifulSoup(html, 'html.parser')
    for script in soup(["script", "style"]):
        script.decompose()
    text = soup.get_text(separator=' ', strip=True)
    return {
        "length_chars": len(text),
        "words": len(re.findall(r'\b\w+\b', text)),
        "sample": text[:500] + "..." if len(text) > 500 else text
    }

def analyze_links(html, base_url):
    """Анализирует ссылки: внутренние, внешние, битые, nofollow"""
    soup = BeautifulSoup(html, 'html.parser')
    all_links = []
    for a in soup.find_all('a', href=True):
        href = a['href'].strip()
        if href.startswith('#') or href.startswith('javascript:') or href.startswith('mailto:'):
            continue
        full_url = urljoin(base_url, href)
        rel = a.get('rel', [])
        nofollow = 'nofollow' in rel
        all_links.append((full_url, nofollow))
    
    unique = list(dict.fromkeys([url for url, _ in all_links]))
    base_domain = urlparse(base_url).netloc
    internal = [url for url in unique if urlparse(url).netloc == base_domain]
    external = [url for url in unique if urlparse(url).netloc != base_domain]
    
    # Проверка битых (первые CHECK_LINKS_LIMIT)
    broken = []
    for link in unique[:CHECK_LINKS_LIMIT]:
        try:
            resp = requests.head(link, timeout=5, allow_redirects=True)
            if resp.status_code >= 400:
                broken.append((link, resp.status_code))
        except Exception:
            broken.append((link, "ошибка соединения"))
    
    # Подсчёт внешних ссылок с nofollow
    external_nofollow = sum(1 for url, nf in all_links if urlparse(url).netloc != base_domain and nf)
    
    return {
        "total_found": len(all_links),
        "unique": len(unique),
        "internal": len(internal),
        "external": len(external),
        "external_nofollow": external_nofollow,
        "broken_checked": len(unique[:CHECK_LINKS_LIMIT]),
        "broken": broken[:10]
    }

def check_duplicates(base_url):
    """Проверяет потенциальные дубли: www, trailing slash, https"""
    variants = []
    parsed = urlparse(base_url)
    # С www
    www_url = f"{parsed.scheme}://www.{parsed.netloc}{parsed.path}"
    if www_url != base_url:
        variants.append(www_url)
    # Без www
    no_www = base_url.replace('www.', '', 1)
    if no_www != base_url:
        variants.append(no_www)
    # Слеш в конце
    if not parsed.path.endswith('/'):
        slash_url = base_url + '/'
        variants.append(slash_url)
    # Без слеша
    if parsed.path.endswith('/') and parsed.path != '/':
        no_slash = base_url.rstrip('/')
        variants.append(no_slash)
    # HTTP/HTTPS
    if parsed.scheme == 'https':
        http_url = 'http://' + parsed.netloc + parsed.path
        variants.append(http_url)
    elif parsed.scheme == 'http':
        https_url = 'https://' + parsed.netloc + parsed.path
        variants.append(https_url)
    
    # Проверяем, какие из вариантов доступны
    accessible = []
    for v in variants:
        r, _, _ = get_html(v)
        if r:
            accessible.append(v)
    return {
        "potential_duplicates": variants,
        "accessible_duplicates": accessible
    }

def check_social(html):
    """Проверяет наличие кнопок поделиться и ссылок на соцсети в футере"""
    soup = BeautifulSoup(html, 'html.parser')
    # Кнопки поделиться (по классам и скриптам)
    share_classes = ['share', 'social', 'soc', 'like']
    buttons = soup.find_all(attrs={'class': lambda c: c and any(p in c.lower() for p in share_classes)})
    share_scripts = soup.find_all('script', src=lambda s: s and ('sharethis' in s or 'addtoany' in s))
    has_share = len(buttons) > 0 or len(share_scripts) > 0
    
    # Ссылки на соцсети в футере
    footer = soup.find('footer')
    social_domains = ['vk.com', 'vkontakte', 'telegram', 'youtube', 'ok.ru', 'facebook', 'twitter', 'instagram']
    footer_links = []
    if footer:
        for a in footer.find_all('a', href=True):
            href = a['href'].lower()
            if any(domain in href for domain in social_domains):
                footer_links.append(href)
    else:
        # если нет footer, ищем по всей странице, но с ограничением
        for a in soup.find_all('a', href=True):
            href = a['href'].lower()
            if any(domain in href for domain in social_domains):
                footer_links.append(href)
    footer_links = list(set(footer_links))[:10]
    
    return {
        "share_buttons": has_share,
        "footer_social_links": footer_links
    }

def check_headers(response):
    """Анализирует HTTP-заголовки: кэширование, gzip"""
    headers = response.headers
    cache_control = headers.get('Cache-Control', '')
    expires = headers.get('Expires', '')
    caching = bool(cache_control or expires)
    gzip = 'gzip' in headers.get('Content-Encoding', '').lower()
    return {
        "caching_enabled": caching,
        "gzip_enabled": gzip
    }

def speed_audit(url):
    """Проверяет скорость загрузки и основные метрики, включая заголовки"""
    try:
        start = time.time()
        resp = requests.get(url, timeout=TIMEOUT)
        load_time = time.time() - start
        ttfb = resp.elapsed.total_seconds()
        size = len(resp.content)
        headers = check_headers(resp)
        return {
            "load_time_sec": round(load_time, 2),
            "ttfb_sec": round(ttfb, 2),
            "size_bytes": size,
            "size_mb": round(size / (1024*1024), 2),
            "headers": headers
        }
    except Exception as e:
        return {"error": str(e)}

def generate_recommendations(data):
    """Формирует список рекомендаций на основе собранных данных"""
    recs = []
    
    # Robots.txt
    if not data['robots']['exists']:
        recs.append("Создайте файл robots.txt и разместите его в корне сайта. Укажите в нём Sitemap.")
    else:
        if not data['robots']['content'] or 'sitemap' not in data['robots']['content'].lower():
            recs.append("В robots.txt добавьте ссылку на sitemap.xml.")
    
    # Мета-теги
    m = data['meta']
    if not m['title'] or m['title_length'] == 0:
        recs.append("Добавьте тег title (уникальный, длиной 50–70 символов).")
    elif m['title_length'] > 70:
        recs.append(f"Сократите title: сейчас {m['title_length']} символов (рекомендуется до 70).")
    if not m['description']:
        recs.append("Добавьте мета-тег description (краткое описание страницы).")
    if not m['keywords']:
        recs.append("Рекомендуется добавить мета-тег keywords (список ключевых слов).")
    if not m['canonical']:
        recs.append("Укажите каноническую ссылку <link rel='canonical'> для избежания дублей.")
    
    # Open Graph
    og = m.get('open_graph', {})
    missing_og = [k for k, v in og.items() if not v]
    if missing_og:
        recs.append(f"Добавьте мета-теги Open Graph для социальных сетей: {', '.join(missing_og)}.")
    
    # Заголовки
    h1_count = len(data['headings']['h1'])
    if h1_count == 0:
        recs.append("На странице отсутствует H1. Добавьте один главный заголовок.")
    elif h1_count > 1:
        recs.append(f"Найдено {h1_count} H1. Оставьте только один.")
    
    # Изображения
    img = data['images']
    if img['total'] > 0:
        if img['missing_alt'] > 0:
            recs.append(f"Заполните атрибут alt для {img['missing_alt']} из {img['total']} изображений.")
        if img['missing_title'] > 0:
            recs.append(f"Добавьте атрибут title для {img['missing_title']} изображений (по желанию).")
        recs.append("Используйте осмысленные имена файлов изображений (латиница, ключевые слова).")
    
    # Текст
    txt = data['text']
    if txt['length_chars'] < 500:
        recs.append(f"Увеличьте объём текста на странице (сейчас {txt['length_chars']} символов, рекомендуется от 500).")
    
    # Ссылки
    l = data['links']
    if l['broken']:
        recs.append(f"Обнаружено {len(l['broken'])} битых ссылок. Исправьте или удалите их.")
    if l['external'] > 0 and l['external_nofollow'] < l['external']:
        recs.append("Для внешних ссылок на сомнительные ресурсы используйте атрибут rel='nofollow'.")
    recs.append("Используйте внутреннюю перелинковку: добавляйте ссылки на связанные страницы, хлебные крошки.")
    
    # Дубли
    dup = data['duplicates']
    if dup['accessible_duplicates']:
        recs.append(f"Найдены доступные зеркала страницы: {', '.join(dup['accessible_duplicates'])}. Настройте 301-редирект на основное зеркало.")
    
    # Скорость и заголовки
    sp = data['speed']
    if 'load_time_sec' in sp:
        if sp['load_time_sec'] > 3:
            recs.append(f"Время загрузки {sp['load_time_sec']} с (желательно < 2 с). Оптимизируйте изображения, включите сжатие, кэширование.")
        if sp.get('size_mb', 0) > 2:
            recs.append(f"Вес страницы {sp['size_mb']} МБ. Сжимайте изображения, минифицируйте код.")
        headers = sp.get('headers', {})
        if not headers.get('gzip_enabled'):
            recs.append("Включите GZIP сжатие на сервере (добавьте правила в .htaccess или настройте в панели хостинга).")
        if not headers.get('caching_enabled'):
            recs.append("Настройте кэширование браузера (Cache-Control, Expires) для статических ресурсов.")
    
    # SMO
    soc = data['social']
    if not soc['share_buttons']:
        recs.append("Добавьте кнопки «Поделиться» (например, AddToAny, ShareThis).")
    if not soc['footer_social_links']:
        recs.append("Укажите ссылки на ваши профили в соцсетях (VK, Telegram и др.) в подвале сайта.")
    
    return recs

def generate_manual_checklist():
    """Возвращает список пунктов, которые требуют ручной оценки."""
    return [
        "Дизайн и юзабилити: проверьте визуальную привлекательность, контрастность, адаптивность, простоту навигации. Меню должно содержать 5–7 пунктов.",
        "Уникальность текста: проверьте с помощью сервисов (Content-watch.ru, Text.ru). Должна быть не менее 95%.",
        "Плотность ключевых слов: оцените, чтобы ключевые фразы встречались равномерно, не было переспама. Оптимально 1–2%.",
        "Перелинковка: убедитесь, что на странице есть хлебные крошки, контекстные ссылки на связанные материалы, рекомендательные блоки.",
        "Исходящие ссылки: проверьте, что ссылки ведут на авторитетные тематические ресурсы, для сомнительных используйте nofollow.",
        "Микроразметка (Schema.org): рекомендуется добавить разметку для организаций, статей, товаров для улучшения сниппетов."
    ]

# ----------------------------------------------------------------------
# Основная функция аудита
# ----------------------------------------------------------------------
def seo_audit(url):
    print(f"Провожу SEO-анализ: {url}")
    report = {
        "url": url,
        "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "robots": {},
        "meta": {},
        "headings": {},
        "images": {},
        "text": {},
        "links": {},
        "duplicates": {},
        "social": {},
        "speed": {},
        "recommendations": [],
        "manual_checklist": generate_manual_checklist()
    }
    
    # Получаем главную страницу
    html, response, load_time = get_html(url)
    if not html:
        report["error"] = "Сайт недоступен"
        return report
    
    # 1. robots.txt
    report["robots"] = check_robots_txt(url)
    
    # 2. Мета-теги (включая Open Graph)
    report["meta"] = analyze_meta(html)
    
    # 3. Заголовки
    report["headings"] = analyze_headings(html)
    
    # 4. Изображения
    report["images"] = analyze_images(html, url)
    
    # 5. Текст
    report["text"] = analyze_text(html)
    
    # 6. Ссылки
    report["links"] = analyze_links(html, url)
    
    # 7. Дубли (зеркала)
    report["duplicates"] = check_duplicates(url)
    
    # 8. Социальные элементы
    report["social"] = check_social(html)
    
    # 9. Скорость и заголовки
    report["speed"] = speed_audit(url)
    
    # 10. Рекомендации (автоматические)
    report["recommendations"] = generate_recommendations(report)
    
    return report

# ----------------------------------------------------------------------
# Сохранение отчётов
# ----------------------------------------------------------------------
def save_json(data, filename):
    with open(filename, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    print(f"JSON отчёт сохранён в {filename}")

def save_docx(data, filename):
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("Для сохранения DOCX установите python-docx: pip install python-docx")
        return
    
    doc = Document()
    heading = doc.add_heading("Отчёт по внутренней SEO-оптимизации сайта", 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    # Основная информация
    doc.add_heading("Общая информация", level=1)
    doc.add_paragraph(f"URL: {data['url']}")
    doc.add_paragraph(f"Дата анализа: {data['timestamp']}")
    if "error" in data:
        doc.add_paragraph(f"Ошибка: {data['error']}")
        doc.save(filename)
        return
    
    # 1. Robots.txt
    doc.add_heading("1. Robots.txt", level=1)
    doc.add_paragraph(f"Наличие: {'Да' if data['robots']['exists'] else 'Нет'}")
    if data['robots']['content']:
        doc.add_paragraph(f"Содержимое:\n{data['robots']['content'][:500]}")
    
    # 2. Мета-теги
    doc.add_heading("2. Мета-теги", level=1)
    m = data['meta']
    doc.add_paragraph(f"Title: {m['title']} (длина: {m['title_length']} симв.)")
    doc.add_paragraph(f"Description: {m['description']}")
    doc.add_paragraph(f"Keywords: {m['keywords']}")
    doc.add_paragraph(f"Robots: {m['robots_meta']}")
    doc.add_paragraph(f"Canonical: {m['canonical']}")
    # Open Graph
    doc.add_heading("Open Graph (для соцсетей)", level=2)
    og = m.get('open_graph', {})
    for k, v in og.items():
        doc.add_paragraph(f"{k}: {v}")
    
    # 3. Заголовки
    doc.add_heading("3. Структура заголовков", level=1)
    for tag, texts in data['headings'].items():
        if texts:
            doc.add_paragraph(f"{tag.upper()}: {', '.join(texts[:3])}")
        else:
            doc.add_paragraph(f"{tag.upper()}: отсутствуют")
    
    # 4. Изображения
    doc.add_heading("4. Изображения", level=1)
    img = data['images']
    doc.add_paragraph(f"Всего изображений: {img['total']}")
    doc.add_paragraph(f"Без alt: {img['missing_alt']}")
    doc.add_paragraph(f"Без title: {img['missing_title']}")
    if img['filenames_examples']:
        doc.add_paragraph(f"Примеры имён файлов: {', '.join(img['filenames_examples'])}")
    
    # 5. Текст
    doc.add_heading("5. Текстовое содержимое", level=1)
    txt = data['text']
    doc.add_paragraph(f"Объём текста: {txt['length_chars']} символов, {txt['words']} слов")
    doc.add_paragraph(f"Фрагмент текста:\n{txt['sample'][:400]}...")
    
    # 6. Ссылки
    doc.add_heading("6. Ссылки", level=1)
    l = data['links']
    doc.add_paragraph(f"Всего ссылок: {l['total_found']} (уникальных: {l['unique']})")
    doc.add_paragraph(f"Внутренних: {l['internal']}, внешних: {l['external']}")
    doc.add_paragraph(f"Из внешних ссылок с nofollow: {l['external_nofollow']}")
    if l['broken']:
        doc.add_paragraph(f"Битых ссылок: {len(l['broken'])}")
        for link, code in l['broken'][:5]:
            doc.add_paragraph(f"  {link} -> {code}")
    else:
        doc.add_paragraph("Битых ссылок не обнаружено (в проверенной выборке).")
    
    # 7. Дубли
    doc.add_heading("7. Потенциальные дубли", level=1)
    dup = data['duplicates']
    if dup['potential_duplicates']:
        doc.add_paragraph("Варианты URL, которые могут дублировать страницу:")
        for v in dup['potential_duplicates']:
            doc.add_paragraph(f"  {v}")
    if dup['accessible_duplicates']:
        doc.add_paragraph("Доступные зеркала (требуется редирект):")
        for a in dup['accessible_duplicates']:
            doc.add_paragraph(f"  {a}")
    
    # 8. Социальные медиа
    doc.add_heading("8. Социальные элементы (SMO)", level=1)
    soc = data['social']
    doc.add_paragraph(f"Кнопки «Поделиться»: {'есть' if soc['share_buttons'] else 'нет'}")
    if soc['footer_social_links']:
        doc.add_paragraph("Ссылки на соцсети в футере:")
        for link in soc['footer_social_links']:
            doc.add_paragraph(f"  {link}")
    else:
        doc.add_paragraph("Ссылки на соцсети в футере отсутствуют.")
    
    # 9. Скорость загрузки и технические параметры
    doc.add_heading("9. Скорость загрузки и настройки сервера", level=1)
    sp = data['speed']
    if 'error' in sp:
        doc.add_paragraph(f"Не удалось измерить: {sp['error']}")
    else:
        doc.add_paragraph(f"Время загрузки: {sp['load_time_sec']} с")
        doc.add_paragraph(f"TTFB: {sp['ttfb_sec']} с")
        doc.add_paragraph(f"Вес страницы: {sp['size_mb']} МБ ({sp['size_bytes']} байт)")
        headers = sp.get('headers', {})
        doc.add_paragraph(f"GZIP сжатие: {'включено' if headers.get('gzip_enabled') else 'отключено'}")
        doc.add_paragraph(f"Кэширование браузера: {'настроено' if headers.get('caching_enabled') else 'не настроено'}")
    
    # 10. Автоматические рекомендации
    doc.add_heading("10. Автоматические рекомендации", level=1)
    for i, rec in enumerate(data['recommendations'], 1):
        doc.add_paragraph(f"{i}. {rec}")
    
    # 11. Ручная проверка (чек-лист)
    doc.add_heading("11. Чек-лист для ручной оценки", level=1)
    for item in data['manual_checklist']:
        doc.add_paragraph(f"• {item}")
    
    doc.save(filename)
    print(f"DOCX отчёт сохранён в {filename}")

# ----------------------------------------------------------------------
# Запуск
# ----------------------------------------------------------------------
if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Внутренний SEO-аудит сайта (расширенная версия)")
    parser.add_argument("url", help="URL сайта для анализа")
    parser.add_argument("--output", "-o", default="seo_audit", help="Базовое имя выходных файлов")
    args = parser.parse_args()
    
    target_url = args.url
    if not target_url.startswith(('http://', 'https://')):
        target_url = 'https://' + target_url
    
    result = seo_audit(target_url)
    
    # Краткий вывод в консоль
    if "error" in result:
        print(f"Ошибка: {result['error']}")
    else:
        print(f"\nАнализ завершён. Найдено {len(result['recommendations'])} автоматических рекомендаций.")
        print(f"Также добавлен чек-лист для ручной оценки ({len(result['manual_checklist'])} пунктов).")
    
    save_json(result, f"{args.output}.json")
    save_docx(result, f"{args.output}.docx")
    print("\nГотово.")