#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Практическая работа №5: Улучшение поведенческих факторов.
Анализирует факторы, влияющие на поведенческие показатели (отказы, время на сайте, глубина).
Сохраняет отчёт в DOCX и JSON.
Использование: python behavior_audit.py <URL> [--output OUTPUT_NAME]
"""

import sys
import requests
from urllib.parse import urlparse, urljoin
from bs4 import BeautifulSoup
import time
import json
import argparse
import re
from datetime import datetime

# Настройки
TIMEOUT = 10
USER_AGENT = "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
CHECK_LINKS_LIMIT = 50

if sys.platform == 'win32':
    sys.stdout.reconfigure(encoding='utf-8')


# ----------------------------------------------------------------------
# Вспомогательные функции
# ----------------------------------------------------------------------
def get_html(url):
    try:
        headers = {'User-Agent': USER_AGENT}
        start = time.time()
        response = requests.get(url, headers=headers, timeout=TIMEOUT, allow_redirects=True)
        load_time = time.time() - start
        if response.encoding is None or response.encoding.lower() == 'iso-8859-1':
            response.encoding = response.apparent_encoding or 'utf-8'
        return response.text, response, load_time
    except Exception as e:
        return None, None, None


def check_theme_clarity(html, url):
    """Проверяет, понятна ли тематика сайта: логотип, слоган, H1"""
    soup = BeautifulSoup(html, 'html.parser')
    # Логотип (ищем img в шапке или элемент с классом logo)
    logo = soup.find('img', alt=re.compile(r'логотип|logo', re.I)) or \
           soup.find('div', class_=re.compile(r'logo', re.I)) or \
           soup.find('a', class_=re.compile(r'logo', re.I))
    # Слоган (часто в теге <p> или <span> рядом с логотипом)
    slogan = soup.find('span', class_=re.compile(r'slogan|tagline', re.I)) or \
             soup.find('p', class_=re.compile(r'slogan|tagline', re.I))
    # Заголовок H1
    h1 = soup.find('h1')
    h1_text = h1.get_text(strip=True) if h1 else None
    
    return {
        "logo_found": logo is not None,
        "slogan_found": slogan is not None,
        "h1_text": h1_text,
        "theme_clear": (logo is not None or h1_text is not None)  # грубая оценка
    }


def check_showcase(html):
    """Проверяет наличие витрины/каталога на главной (товары, карточки)"""
    soup = BeautifulSoup(html, 'html.parser')
    # Ищем типичные классы для товаров
    products = soup.find_all(class_=re.compile(r'product|item|card|good', re.I))
    # Также смотрим на наличие каталога в меню
    catalog_link = soup.find('a', href=re.compile(r'catalog|katalog|shop', re.I))
    return {
        "has_products": len(products) > 0,
        "catalog_link": catalog_link is not None,
        "product_count": len(products)
    }


def check_problem_solver(html):
    """Проверяет наличие решателя задач: калькулятор, форма поиска, фильтры"""
    soup = BeautifulSoup(html, 'html.parser')
    # Калькулятор
    calculator = soup.find(class_=re.compile(r'calc|calculator', re.I)) or \
                 soup.find('input', {'name': re.compile(r'calc', re.I)})
    # Поиск
    search = soup.find('form', {'role': 'search'}) or \
             soup.find('input', {'type': 'search'}) or \
             soup.find('input', {'name': 's'}) or \
             soup.find('button', {'type': 'submit'}, text=re.compile(r'найти', re.I))
    # Фильтры (для интернет-магазинов)
    filters = soup.find(class_=re.compile(r'filter', re.I))
    return {
        "has_calculator": calculator is not None,
        "has_search": search is not None,
        "has_filters": filters is not None
    }


def check_navigation(html, url):
    """Проверяет качество навигации: меню, активный пункт, хлебные крошки"""
    soup = BeautifulSoup(html, 'html.parser')
    # Главное меню (обычно nav или ul с классами menu)
    menu = soup.find('nav') or soup.find('ul', class_=re.compile(r'menu|nav', re.I))
    # Активный пункт меню
    active = soup.find('li', class_=re.compile(r'active|current', re.I)) or \
             soup.find('a', class_=re.compile(r'active|current', re.I))
    # Хлебные крошки
    breadcrumbs = soup.find('div', class_=re.compile(r'breadcrumb', re.I)) or \
                  soup.find('ul', class_=re.compile(r'breadcrumb', re.I))
    return {
        "has_menu": menu is not None,
        "has_active_item": active is not None,
        "has_breadcrumbs": breadcrumbs is not None
    }


def check_interactivity(html):
    """Проверяет наличие интерактивных элементов: отзывы, голосования, чат, онлайн-консультант"""
    soup = BeautifulSoup(html, 'html.parser')
    reviews = soup.find(class_=re.compile(r'review|comment|otziv', re.I))
    voting = soup.find(class_=re.compile(r'vote|poll|golos', re.I))
    chat = soup.find('div', id=re.compile(r'chat', re.I)) or \
           soup.find('script', src=re.compile(r'jivo|tawk|livechat', re.I))
    online_consultant = soup.find('a', href=re.compile(r'online|consult', re.I)) or \
                        soup.find('div', class_=re.compile(r'consultant', re.I))
    return {
        "has_reviews": reviews is not None,
        "has_voting": voting is not None,
        "has_chat": chat is not None,
        "has_online_consultant": online_consultant is not None
    }


def check_articles(html):
    """Проверяет наличие тематических статей / блога"""
    soup = BeautifulSoup(html, 'html.parser')
    blog_link = soup.find('a', href=re.compile(r'blog|articles|news', re.I))
    articles = soup.find_all('article') or soup.find_all(class_=re.compile(r'article|post|news', re.I))
    return {
        "has_blog_section": blog_link is not None,
        "article_count": len(articles)
    }


def analyze_internal_links(html, base_url):
    """Анализирует внутреннюю перелинковку (количество внутренних ссылок)"""
    soup = BeautifulSoup(html, 'html.parser')
    all_links = []
    for a in soup.find_all('a', href=True):
        href = a['href'].strip()
        if href.startswith('#') or href.startswith('javascript:') or href.startswith('mailto:'):
            continue
        full_url = urljoin(base_url, href)
        all_links.append(full_url)
    
    unique = list(dict.fromkeys(all_links))
    base_domain = urlparse(base_url).netloc
    internal = [l for l in unique if urlparse(l).netloc == base_domain]
    return {
        "total_links": len(all_links),
        "unique_links": len(unique),
        "internal_links": len(internal)
    }


def check_external_links_behavior(html, base_url):
    """Проверяет, открываются ли внешние ссылки в новой вкладке (target='_blank')"""
    soup = BeautifulSoup(html, 'html.parser')
    base_domain = urlparse(base_url).netloc
    external_links = []
    for a in soup.find_all('a', href=True):
        href = a['href'].strip()
        if href.startswith('#') or href.startswith('javascript:') or href.startswith('mailto:'):
            continue
        full_url = urljoin(base_url, href)
        if urlparse(full_url).netloc != base_domain:
            target = a.get('target', '')
            external_links.append({
                "url": full_url,
                "target_blank": target == '_blank'
            })
    return {
        "external_links_count": len(external_links),
        "open_in_new_tab_count": sum(1 for el in external_links if el['target_blank']),
        "examples": external_links[:5]
    }


def measure_speed(url):
    """Измеряет скорость загрузки (время полной загрузки)"""
    try:
        start = time.time()
        resp = requests.get(url, timeout=TIMEOUT, allow_redirects=True)
        load_time = time.time() - start
        return round(load_time, 2)
    except Exception:
        return None


def check_404_page(base_url):
    """Проверяет наличие оформленной страницы 404"""
    # Генерируем заведомо несуществующий URL
    test_url = urljoin(base_url, "/nonexistent_page_" + str(int(time.time())))
    try:
        resp = requests.get(test_url, timeout=5, allow_redirects=False)
        # Если статус 404 и страница не пустая
        if resp.status_code == 404 and len(resp.content) > 500:
            return {"exists": True, "status_code": 404}
        else:
            return {"exists": False, "status_code": resp.status_code}
    except Exception:
        return {"exists": False, "error": "не удалось проверить"}


def generate_recommendations(data):
    """Формирует рекомендации на основе собранных данных"""
    recs = []
    
    # Тематика
    theme = data['theme_clarity']
    if not theme['logo_found']:
        recs.append("Добавьте логотип, чтобы посетитель сразу понимал, на каком сайте находится.")
    if not theme['h1_text']:
        recs.append("Разместите заголовок H1, отражающий тематику сайта.")
    if theme['h1_text'] and len(theme['h1_text']) < 10:
        recs.append("Сделайте заголовок H1 более информативным.")
    
    # Витрина
    showcase = data['showcase']
    if not showcase['has_products'] and not showcase['catalog_link']:
        recs.append("На главной странице разместите витрину товаров или ссылку на каталог.")
    
    # Решение задач
    solver = data['problem_solver']
    if not solver['has_calculator'] and not solver['has_search'] and not solver['has_filters']:
        recs.append("Добавьте инструмент, решающий конкретную задачу посетителя (калькулятор, умный поиск, фильтры).")
    
    # Навигация
    nav = data['navigation']
    if not nav['has_menu']:
        recs.append("Создайте главное меню для удобной навигации.")
    if not nav['has_active_item']:
        recs.append("Подсвечивайте активный пункт меню, чтобы посетитель понимал, где находится.")
    if not nav['has_breadcrumbs']:
        recs.append("Добавьте хлебные крошки для облегчения навигации.")
    
    # Интерактив
    inter = data['interactivity']
    if not inter['has_reviews']:
        recs.append("Добавьте систему отзывов — это повышает доверие и удерживает пользователей.")
    if not inter['has_chat'] and not inter['has_online_consultant']:
        recs.append("Рассмотрите возможность установки онлайн-консультанта или чата.")
    
    # Статьи
    articles = data['articles']
    if not articles['has_blog_section'] and articles['article_count'] == 0:
        recs.append("Создайте раздел с тематическими статьями или новостями — это увеличит время на сайте.")
    
    # Перелинковка
    links = data['internal_links']
    if links['internal_links'] < 3:
        recs.append("Увеличьте количество внутренних ссылок, связывайте страницы между собой.")
    
    # Внешние ссылки
    ext = data['external_links_behavior']
    if ext['external_links_count'] > 0:
        if ext['open_in_new_tab_count'] < ext['external_links_count']:
            recs.append("Настройте открытие внешних ссылок в новой вкладке (target='_blank'), чтобы посетитель не покидал ваш сайт.")
    
    # Скорость
    speed = data['speed_sec']
    if speed and speed > 3:
        recs.append(f"Скорость загрузки {speed} с. Оптимизируйте загрузку (желательно < 2 с).")
    
    # Страница 404
    p404 = data['404_page']
    if not p404.get('exists'):
        recs.append("Создайте оформленную страницу 404 с навигацией и ссылками на важные разделы.")
    
    # Поведенческие факторы (ручной ввод)
    recs.append("Заполните раздел 'Поведенческие факторы' в отчёте реальными данными из Яндекс.Метрики или Google Analytics.")
    
    return recs


def generate_manual_sections():
    """Возвращает шаблон для ручного заполнения поведенческих факторов"""
    return {
        "behavioral_data": {
            "bounce_rate_percent": "",
            "time_on_site_sec": "",
            "page_depth": "",
            "return_visits": "",
            "analytics_tools": "",
            "notes": ""
        },
        "return_mechanisms": {
            "stock_notify": "",
            "subscription": "",
            "bookmarks": "",
            "other": ""
        },
        "expert_comments": ""
    }


def behavioral_audit(url):
    print(f"Анализ поведенческих факторов: {url}")
    report = {
        "url": url,
        "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "theme_clarity": {},
        "showcase": {},
        "problem_solver": {},
        "navigation": {},
        "interactivity": {},
        "articles": {},
        "internal_links": {},
        "external_links_behavior": {},
        "speed_sec": None,
        "404_page": {},
        "manual": generate_manual_sections(),
        "recommendations": []
    }
    
    html, response, load_time = get_html(url)
    if not html:
        report["error"] = "Сайт недоступен"
        return report
    
    report["speed_sec"] = round(load_time, 2)
    report["theme_clarity"] = check_theme_clarity(html, url)
    report["showcase"] = check_showcase(html)
    report["problem_solver"] = check_problem_solver(html)
    report["navigation"] = check_navigation(html, url)
    report["interactivity"] = check_interactivity(html)
    report["articles"] = check_articles(html)
    report["internal_links"] = analyze_internal_links(html, url)
    report["external_links_behavior"] = check_external_links_behavior(html, url)
    report["404_page"] = check_404_page(url)
    report["recommendations"] = generate_recommendations(report)
    
    return report


def save_json(data, filename):
    with open(filename, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    print(f"JSON отчёт сохранён в {filename}")


def save_docx(data, filename):
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("Для сохранения DOCX установите python-docx: pip install python-docx")
        return
    
    doc = Document()
    heading = doc.add_heading("Отчёт по улучшению поведенческих факторов", 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    doc.add_paragraph(f"URL: {data['url']}")
    doc.add_paragraph(f"Дата анализа: {data['timestamp']}")
    doc.add_paragraph()
    
    if "error" in data:
        doc.add_paragraph(f"Ошибка: {data['error']}")
        doc.save(filename)
        return
    
    # 1. Тематика
    doc.add_heading("1. Понятность тематики", level=1)
    t = data['theme_clarity']
    doc.add_paragraph(f"Логотип: {'есть' if t['logo_found'] else 'нет'}")
    doc.add_paragraph(f"Слоган: {'есть' if t['slogan_found'] else 'нет'}")
    doc.add_paragraph(f"Заголовок H1: {t['h1_text']}")
    doc.add_paragraph(f"Тематика понятна: {'да' if t['theme_clear'] else 'нет (рекомендуется доработать)'}")
    
    # 2. Витрина
    doc.add_heading("2. Витрина / каталог на главной", level=1)
    s = data['showcase']
    doc.add_paragraph(f"Товары/карточки на главной: {'есть' if s['has_products'] else 'нет'}")
    doc.add_paragraph(f"Ссылка на каталог: {'есть' if s['catalog_link'] else 'нет'}")
    if s['has_products']:
        doc.add_paragraph(f"Количество товарных элементов: {s['product_count']}")
    
    # 3. Решение задач
    doc.add_heading("3. Инструменты решения задач", level=1)
    p = data['problem_solver']
    doc.add_paragraph(f"Калькулятор: {'есть' if p['has_calculator'] else 'нет'}")
    doc.add_paragraph(f"Поиск: {'есть' if p['has_search'] else 'нет'}")
    doc.add_paragraph(f"Фильтры: {'есть' if p['has_filters'] else 'нет'}")
    
    # 4. Навигация
    doc.add_heading("4. Качество навигации", level=1)
    n = data['navigation']
    doc.add_paragraph(f"Главное меню: {'есть' if n['has_menu'] else 'нет'}")
    doc.add_paragraph(f"Активный пункт меню: {'подсвечен' if n['has_active_item'] else 'не подсвечен'}")
    doc.add_paragraph(f"Хлебные крошки: {'есть' if n['has_breadcrumbs'] else 'нет'}")
    
    # 5. Интерактив
    doc.add_heading("5. Интерактивные элементы", level=1)
    i = data['interactivity']
    doc.add_paragraph(f"Отзывы: {'есть' if i['has_reviews'] else 'нет'}")
    doc.add_paragraph(f"Голосования/опросы: {'есть' if i['has_voting'] else 'нет'}")
    doc.add_paragraph(f"Онлайн-чат: {'есть' if i['has_chat'] else 'нет'}")
    doc.add_paragraph(f"Онлайн-консультант: {'есть' if i['has_online_consultant'] else 'нет'}")
    
    # 6. Статьи
    doc.add_heading("6. Тематические статьи / блог", level=1)
    a = data['articles']
    doc.add_paragraph(f"Раздел статей/блога: {'есть' if a['has_blog_section'] else 'нет'}")
    doc.add_paragraph(f"Количество статей на главной: {a['article_count']}")
    
    # 7. Перелинковка
    doc.add_heading("7. Внутренняя перелинковка", level=1)
    l = data['internal_links']
    doc.add_paragraph(f"Всего ссылок на странице: {l['total_links']}")
    doc.add_paragraph(f"Уникальных ссылок: {l['unique_links']}")
    doc.add_paragraph(f"Внутренних ссылок: {l['internal_links']}")
    
    # 8. Внешние ссылки
    doc.add_heading("8. Поведение внешних ссылок", level=1)
    e = data['external_links_behavior']
    doc.add_paragraph(f"Количество внешних ссылок: {e['external_links_count']}")
    doc.add_paragraph(f"Из них открываются в новой вкладке: {e['open_in_new_tab_count']}")
    if e['examples']:
        doc.add_paragraph("Примеры:")
        for ex in e['examples'][:3]:
            doc.add_paragraph(f"  {ex['url']} -> target='_blank': {ex['target_blank']}")
    
    # 9. Скорость
    doc.add_heading("9. Скорость загрузки", level=1)
    doc.add_paragraph(f"Время полной загрузки главной страницы: {data['speed_sec']} с")
    
    # 10. Страница 404
    doc.add_heading("10. Страница 404", level=1)
    p404 = data['404_page']
    if p404.get('exists'):
        doc.add_paragraph("Страница 404: присутствует (оформленная)")
    else:
        doc.add_paragraph("Страница 404: отсутствует или не оформлена")
    
    # 11. Поведенческие факторы (ручной ввод)
    doc.add_heading("11. Поведенческие факторы (заполнить из аналитики)", level=1)
    m = data['manual']['behavioral_data']
    doc.add_paragraph("Показатель отказов (%):")
    doc.add_paragraph(m['bounce_rate_percent'] or "_____________________")
    doc.add_paragraph("Среднее время на сайте (сек):")
    doc.add_paragraph(m['time_on_site_sec'] or "_____________________")
    doc.add_paragraph("Глубина просмотра (кол-во страниц):")
    doc.add_paragraph(m['page_depth'] or "_____________________")
    doc.add_paragraph("Возвраты посетителей (%):")
    doc.add_paragraph(m['return_visits'] or "_____________________")
    doc.add_paragraph("Используемые системы аналитики:")
    doc.add_paragraph(m['analytics_tools'] or "_____________________")
    doc.add_paragraph("Примечания:")
    doc.add_paragraph(m['notes'] or "_____________________")
    
    # 12. Механизмы возврата
    doc.add_heading("12. Механизмы возвращения посетителей", level=1)
    ret = data['manual']['return_mechanisms']
    doc.add_paragraph("Сообщить о поступлении товара:")
    doc.add_paragraph(ret['stock_notify'] or "_____________________")
    doc.add_paragraph("Подписка на рассылку:")
    doc.add_paragraph(ret['subscription'] or "_____________________")
    doc.add_paragraph("Возможность добавить в закладки:")
    doc.add_paragraph(ret['bookmarks'] or "_____________________")
    doc.add_paragraph("Другие способы:")
    doc.add_paragraph(ret['other'] or "_____________________")
    
    # 13. Комментарии эксперта
    doc.add_heading("13. Комментарии эксперта", level=1)
    doc.add_paragraph(data['manual']['expert_comments'] or "_____________________")
    
    # 14. Рекомендации
    doc.add_heading("14. Рекомендации по улучшению поведенческих факторов", level=1)
    for i, rec in enumerate(data['recommendations'], 1):
        doc.add_paragraph(f"{i}. {rec}")
    
    doc.save(filename)
    print(f"DOCX отчёт сохранён в {filename}")


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Аудит поведенческих факторов")
    parser.add_argument("url", help="URL сайта для анализа")
    parser.add_argument("--output", "-o", default="behavior_audit", help="Базовое имя выходных файлов")
    args = parser.parse_args()
    
    target_url = args.url
    if not target_url.startswith(('http://', 'https://')):
        target_url = 'https://' + target_url
    
    result = behavioral_audit(target_url)
    
    if "error" in result:
        print(f"Ошибка: {result['error']}")
    else:
        print(f"\nАнализ завершён. Найдено {len(result['recommendations'])} рекомендаций.")
    
    save_json(result, f"{args.output}.json")
    save_docx(result, f"{args.output}.docx")
    print("\nГотово.")