#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
SEO и юзабилити аудит сайта.
Сохраняет результаты в DOCX и JSON.
Использование: python audit_full.py <URL> [--output OUTPUT_NAME]
"""

import sys
import requests
from urllib.parse import urlparse, urljoin
from bs4 import BeautifulSoup
import time
import re
import json
import argparse
import html  # для экранирования при работе с Word (на всякий случай)
from datetime import datetime

# Настройки
TIMEOUT = 10
USER_AGENT = "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36"
CHECK_LINKS_LIMIT = 30

# Для корректного вывода в консоль (Windows)
if sys.platform == 'win32':
    sys.stdout.reconfigure(encoding='utf-8')

# ----------------------------------------------------------------------
# Функции для загрузки данных
# ----------------------------------------------------------------------
def get_html(url):
    """Получить HTML страницы с обработкой ошибок и правильной кодировкой"""
    try:
        headers = {'User-Agent': USER_AGENT}
        start = time.time()
        response = requests.get(url, headers=headers, timeout=TIMEOUT, allow_redirects=True)
        load_time = time.time() - start
        # Устанавливаем кодировку по содержимому
        if response.encoding is None or response.encoding.lower() == 'iso-8859-1':
            response.encoding = response.apparent_encoding or 'utf-8'
        if response.status_code == 200:
            return response.text, load_time, response.url
        else:
            return None, load_time, response.url
    except Exception as e:
        print(f"Ошибка загрузки {url}: {e}")
        return None, None, None

def check_robots_txt(base_url):
    """Проверка наличия и содержимого robots.txt"""
    robots_url = urljoin(base_url, "/robots.txt")
    html, load_time, final_url = get_html(robots_url)
    if html:
        return {"exists": True, "content": html[:500] + "..." if len(html) > 500 else html}
    else:
        return {"exists": False, "content": None}

def check_sitemap(base_url):
    """Проверка наличия sitemap.xml (по умолчанию)"""
    sitemap_url = urljoin(base_url, "/sitemap.xml")
    html, load_time, final_url = get_html(sitemap_url)
    return {"exists": html is not None}

def analyze_meta(html, url):
    """Анализ метатегов: title, description, keywords, robots, canonical"""
    soup = BeautifulSoup(html, 'html.parser')
    result = {}
    
    # Title
    title_tag = soup.find('title')
    result['title'] = title_tag.get_text(strip=True) if title_tag else None
    result['title_length'] = len(result['title']) if result['title'] else 0
    
    # Meta description
    meta_desc = soup.find('meta', attrs={'name': 'description'})
    result['description'] = meta_desc.get('content') if meta_desc else None
    
    # Meta keywords
    meta_keys = soup.find('meta', attrs={'name': 'keywords'})
    result['keywords'] = meta_keys.get('content') if meta_keys else None
    
    # Meta robots
    meta_robots = soup.find('meta', attrs={'name': 'robots'})
    result['robots_meta'] = meta_robots.get('content') if meta_robots else None
    
    # Canonical
    canonical = soup.find('link', attrs={'rel': 'canonical'})
    result['canonical'] = canonical.get('href') if canonical else None
    
    return result

def analyze_headings(html):
    """Анализ заголовков H1-H6"""
    soup = BeautifulSoup(html, 'html.parser')
    headings = {}
    for i in range(1, 7):
        tag = f'h{i}'
        elements = soup.find_all(tag)
        headings[tag] = [elem.get_text(strip=True) for elem in elements]
    return headings

def analyze_images(html, base_url):
    """Анализ изображений: количество, наличие alt и title"""
    soup = BeautifulSoup(html, 'html.parser')
    images = soup.find_all('img')
    total = len(images)
    missing_alt = sum(1 for img in images if not img.get('alt'))
    missing_title = sum(1 for img in images if not img.get('title'))
    return {
        "total_images": total,
        "missing_alt": missing_alt,
        "missing_title": missing_title,
        "alt_percentage": 100 - (missing_alt / total * 100) if total else 100,
        "title_percentage": 100 - (missing_title / total * 100) if total else 100
    }

def analyze_text(html):
    """Анализ текста: объём, примерная плотность ключевых слов (заглушка)"""
    soup = BeautifulSoup(html, 'html.parser')
    # Удаляем скрипты и стили
    for script in soup(["script", "style"]):
        script.decompose()
    text = soup.get_text(separator=' ', strip=True)
    text_length = len(text)
    words = re.findall(r'\b\w+\b', text.lower())
    word_count = len(words)
    return {
        "text_length_chars": text_length,
        "word_count": word_count,
        "sample_text": text[:500] + "..." if text_length > 500 else text
    }

def check_links(html, base_url):
    """Проверка внутренних и внешних ссылок, поиск битых (первые CHECK_LINKS_LIMIT)"""
    soup = BeautifulSoup(html, 'html.parser')
    all_links = []
    for a in soup.find_all('a', href=True):
        href = a['href'].strip()
        if href.startswith('#') or href.startswith('javascript:') or href.startswith('mailto:') or href.startswith('tel:'):
            continue
        full_url = urljoin(base_url, href)
        all_links.append(full_url)
    
    unique_links = list(dict.fromkeys(all_links))
    links_to_check = unique_links[:CHECK_LINKS_LIMIT]
    
    broken = []
    for link in links_to_check:
        try:
            resp = requests.head(link, timeout=5, allow_redirects=True)
            if resp.status_code >= 400:
                broken.append((link, resp.status_code))
        except Exception:
            broken.append((link, "ошибка соединения"))
    
    base_domain = urlparse(base_url).netloc
    internal = [l for l in unique_links if urlparse(l).netloc == base_domain]
    external = [l for l in unique_links if urlparse(l).netloc != base_domain]
    
    return {
        "total_links_found": len(all_links),
        "unique_links": len(unique_links),
        "internal_links": len(internal),
        "external_links": len(external),
        "broken_links_checked": len(links_to_check),
        "broken_links": broken[:10]
    }

def measure_speed(url):
    """Измерение скорости загрузки главной страницы (среднее из 3 попыток)"""
    times = []
    for _ in range(3):
        try:
            start = time.time()
            requests.get(url, timeout=TIMEOUT, allow_redirects=True)
            times.append(time.time() - start)
        except:
            times.append(None)
    valid = [t for t in times if t is not None]
    avg = sum(valid) / len(valid) if valid else None
    return avg

def check_social_buttons(html):
    """Проверка наличия кнопок «Поделиться» по типичным классам"""
    soup = BeautifulSoup(html, 'html.parser')
    share_patterns = ['share', 'social', 'soc', 'like', 'twitter', 'facebook', 'vkontakte', 'telegram', 'whatsapp']
    buttons = soup.find_all(attrs={'class': lambda c: c and any(p in c.lower() for p in share_patterns)})
    # Также проверим iframe или скрипты виджетов
    widgets = soup.find_all('script', src=lambda s: s and ('sharethis' in s or 'addtoany' in s))
    return len(buttons) > 0 or len(widgets) > 0

def check_social_footer_links(html, base_url):
    """Проверка наличия ссылок на соцсети в подвале (footer)"""
    soup = BeautifulSoup(html, 'html.parser')
    footer = soup.find('footer')
    if not footer:
        # если нет footer, ищем по всей странице
        footer = soup
    social_domains = ['vk.com', 'vkontakte', 'telegram', 'youtube', 'ok.ru', 'facebook', 'twitter', 'instagram']
    links = footer.find_all('a', href=True)
    found = []
    for a in links:
        href = a['href'].lower()
        if any(domain in href for domain in social_domains):
            found.append(href)
    return list(set(found))

# ----------------------------------------------------------------------
# Сохранение в файлы
# ----------------------------------------------------------------------
def save_json(report_dict, filename):
    with open(filename, 'w', encoding='utf-8') as f:
        json.dump(report_dict, f, ensure_ascii=False, indent=2)
    print(f"JSON отчёт сохранён в {filename}")

def save_docx(report_text, filename):
    """Сохраняет отчёт в DOCX с помощью python-docx"""
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("Для сохранения DOCX установите python-docx: pip install python-docx")
        return
    
    doc = Document()
    heading = doc.add_heading("Отчёт по аудиту сайта", 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    for line in report_text.split('\n'):
        if line.strip():
            # Для DOCX не нужно экранировать HTML, просто добавляем строку
            if line.startswith('=') or line.startswith('-'):
                doc.add_paragraph(line)
            elif line.startswith('    '):
                p = doc.add_paragraph(line, style='Normal')
                p.paragraph_format.left_indent = Pt(20)
            else:
                p = doc.add_paragraph(line, style='Normal')
    doc.save(filename)
    print(f"DOCX отчёт сохранён в {filename}")

# ----------------------------------------------------------------------
# Основная функция аудита
# ----------------------------------------------------------------------
def audit_site(url):
    print(f"Начинаю аудит сайта: {url}")
    report = {
        "url": url,
        "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "availability": {},
        "robots": {},
        "sitemap": {},
        "meta": {},
        "headings": {},
        "images": {},
        "text": {},
        "links": {},
        "speed": {},
        "social": {}
    }
    
    # 1. Проверка доступности
    html, load_time, final_url = get_html(url)
    if not html:
        report["availability"] = {"accessible": False, "error": "Сайт не загружен или вернул ошибку"}
        print("Сайт недоступен, дальнейший анализ невозможен.")
        return report
    report["availability"] = {"accessible": True, "load_time_sec": round(load_time, 2), "final_url": final_url}
    
    # 2. robots.txt
    report["robots"] = check_robots_txt(url)
    
    # 3. sitemap.xml
    report["sitemap"] = check_sitemap(url)
    
    # 4. Метатеги
    report["meta"] = analyze_meta(html, url)
    
    # 5. Заголовки
    report["headings"] = analyze_headings(html)
    
    # 6. Изображения
    report["images"] = analyze_images(html, url)
    
    # 7. Текст
    report["text"] = analyze_text(html)
    
    # 8. Ссылки
    report["links"] = check_links(html, url)
    
    # 9. Скорость
    speed = measure_speed(url)
    report["speed"] = {"avg_load_time_3": round(speed, 2) if speed else None}
    
    # 10. Социальные элементы
    report["social"] = {
        "share_buttons": check_social_buttons(html),
        "footer_social_links": check_social_footer_links(html, url)
    }
    
    return report

def generate_report_text(report):
    """Генерирует текст отчёта для вывода и сохранения"""
    lines = []
    lines.append("="*80)
    lines.append(f"ОТЧЁТ ПО АУДИТУ САЙТА: {report['url']}")
    lines.append("="*80)
    
    av = report['availability']
    if not av.get('accessible'):
        lines.append(f"\n[ОШИБКА] Сайт недоступен: {av.get('error')}")
        return "\n".join(lines)
    
    lines.append(f"\n1. ДОСТУПНОСТЬ: OK (время загрузки: {av['load_time_sec']} с, финальный URL: {av['final_url']})")
    
    r = report['robots']
    lines.append(f"\n2. ROBOTS.TXT: {'присутствует' if r['exists'] else 'отсутствует'}")
    if r['exists']:
        lines.append(f"   Содержимое: {r['content'][:200]}...")
    
    s = report['sitemap']
    lines.append(f"\n3. SITEMAP.XML: {'присутствует' if s['exists'] else 'отсутствует'}")
    
    m = report['meta']
    lines.append("\n4. МЕТАТЕГИ:")
    lines.append(f"   Title: {m['title']} (длина: {m['title_length']} симв.)")
    lines.append(f"   Description: {m['description']}")
    lines.append(f"   Keywords: {m['keywords']}")
    lines.append(f"   Robots: {m['robots_meta']}")
    lines.append(f"   Canonical: {m['canonical']}")
    
    h = report['headings']
    lines.append("\n5. ЗАГОЛОВКИ:")
    for tag, texts in h.items():
        if texts:
            lines.append(f"   {tag.upper()}: {texts}")
        else:
            lines.append(f"   {tag.upper()}: не найдены")
    
    img = report['images']
    lines.append("\n6. ИЗОБРАЖЕНИЯ:")
    lines.append(f"   Всего: {img['total_images']}")
    lines.append(f"   Без alt: {img['missing_alt']} ({100 - img['alt_percentage']:.1f}%)")
    lines.append(f"   Без title: {img['missing_title']} ({100 - img['title_percentage']:.1f}%)")
    
    txt = report['text']
    lines.append("\n7. ТЕКСТОВОЕ СОДЕРЖАНИЕ:")
    lines.append(f"   Объём текста (символов): {txt['text_length_chars']}")
    lines.append(f"   Количество слов: {txt['word_count']}")
    lines.append(f"   Пример текста: {txt['sample_text'][:200]}...")
    
    l = report['links']
    lines.append("\n8. ССЫЛКИ:")
    lines.append(f"   Всего ссылок: {l['total_links_found']} (уникальных: {l['unique_links']})")
    lines.append(f"   Внутренних: {l['internal_links']}, внешних: {l['external_links']}")
    if l['broken_links_checked']:
        lines.append(f"   Битых ссылок (из {l['broken_links_checked']} проверенных): {len(l['broken_links'])}")
        if l['broken_links']:
            lines.append("   Примеры битых ссылок:")
            for link, code in l['broken_links'][:5]:
                lines.append(f"      {link} -> {code}")
    
    sp = report['speed']
    lines.append(f"\n9. СКОРОСТЬ ЗАГРУЗКИ (среднее 3 замеров): {sp['avg_load_time_3']} с")
    
    soc = report['social']
    lines.append("\n10. СОЦИАЛЬНЫЕ МЕДИА (SMO):")
    lines.append(f"    Кнопки «Поделиться»: {'присутствуют' if soc['share_buttons'] else 'отсутствуют'}")
    if soc['footer_social_links']:
        lines.append(f"    Ссылки на соцсети в футере: {', '.join(soc['footer_social_links'])}")
    else:
        lines.append("    Ссылки на соцсети в футере: отсутствуют")
    
    lines.append("\n11. ПЕРВООЧЕРЕДНЫЕ РЕКОМЕНДАЦИИ:")
    recs = []
    
    # Title
    if m['title_length'] == 0:
        recs.append("- Отсутствует тег Title – добавьте уникальный заголовок (50–70 символов)")
    elif m['title_length'] > 70:
        recs.append(f"- Длина Title ({m['title_length']} симв.) превышает рекомендуемую (до 70)")
    
    # Description
    if not m['description']:
        recs.append("- Добавьте мета-тег description с уникальным описанием страницы")
    
    # Keywords
    if not m['keywords']:
        recs.append("- Рекомендуется добавить мета-тег keywords (список ключевых слов)")
    
    # Canonical
    if not m['canonical']:
        recs.append("- Добавьте каноническую ссылку <link rel=\"canonical\"> для избежания дублей")
    
    # H1
    h1_count = len(h.get('h1', []))
    if h1_count == 0:
        recs.append("- На странице отсутствует главный заголовок H1")
    elif h1_count > 1:
        recs.append(f"- На странице найдено {h1_count} H1, оставьте только один")
    
    # Изображения
    if img['missing_alt'] > 0:
        recs.append(f"- Заполните атрибут alt для {img['missing_alt']} изображений")
    
    # Текст
    if txt['text_length_chars'] < 500:
        recs.append(f"- Увеличьте объём текста на странице (сейчас {txt['text_length_chars']} символов, рекомендуется от 500)")
    
    # Битые ссылки
    if l['broken_links']:
        recs.append("- Устраните битые ссылки")
    
    # Скорость
    if sp['avg_load_time_3'] and sp['avg_load_time_3'] > 3:
        recs.append(f"- Увеличьте скорость загрузки (сейчас {sp['avg_load_time_3']} с, желательно < 2 с)")
    
    # robots.txt
    if not r['exists']:
        recs.append("- Создайте robots.txt и настройте правила индексации")
    
    # sitemap.xml
    if not s['exists']:
        recs.append("- Добавьте sitemap.xml и укажите его в robots.txt")
    
    # Социальные кнопки
    if not soc['share_buttons']:
        recs.append("- Добавьте кнопки «Поделиться» (например, AddToAny, ShareThis)")
    
    if not soc['footer_social_links']:
        recs.append("- Укажите ссылки на ваши профили в социальных сетях в подвале сайта")
    
    if recs:
        for rec in recs:
            lines.append(rec)
    else:
        lines.append("- Основных проблем не выявлено, требуется более детальный анализ.")
    
    lines.append("\n" + "="*80)
    return "\n".join(lines)

# ----------------------------------------------------------------------
# Запуск
# ----------------------------------------------------------------------
if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="SEO и юзабилити аудит сайта")
    parser.add_argument("url", help="URL сайта для аудита")
    parser.add_argument("--output", "-o", default="audit", help="Базовое имя выходных файлов (без расширения)")
    args = parser.parse_args()
    
    target_url = args.url
    if not target_url.startswith(('http://', 'https://')):
        target_url = 'https://' + target_url
    
    result = audit_site(target_url)
    report_text = generate_report_text(result)
    
    # Вывод в консоль
    print(report_text)
    
    # Сохранение в JSON
    json_file = f"{args.output}.json"
    save_json(result, json_file)
    
    # Сохранение в DOCX
    docx_file = f"{args.output}.docx"
    save_docx(report_text, docx_file)
    
    print("\nГотово.")